"""Generate Presentation DOCX — Hospital Claim AI สำหรับ present ผอ.รพ.พญาไทศรีราชา"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_thai_font(run, size=14, bold=False, color=None):
    run.font.name = "TH SarabunPSK"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:cs"), "TH SarabunPSK")


def add_p(doc, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_thai_font(run, size, bold, color)
    return p


def add_bullet(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.style = "List Bullet"
    run = p.add_run(text)
    set_thai_font(run, size, bold)
    return p


def add_number(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.style = "List Number"
    run = p.add_run(text)
    set_thai_font(run, size, bold)
    return p


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_thai_font(run, 14, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, header_color)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            set_thai_font(run, 14)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def generate_presentation(output_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════
    add_p(doc, "", 24)
    add_p(doc, "", 24)
    add_p(doc, "", 24)
    add_p(doc, "Hospital Claim AI", 36, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "ระบบ AI ช่วยลด Deny Rate และเพิ่มรายได้จากการเบิก สปสช.", 20, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x44, 0x72, 0xC4))
    add_p(doc, "", 14)
    add_p(doc, "โรงพยาบาลพญาไทศรีราชา", 22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "รหัสหน่วยบริการ 11855 | เขต สปสช. 6 ระยอง", 16, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x66, 0x66, 0x66))
    add_p(doc, "", 14)
    add_p(doc, "", 14)
    add_p(doc, "มีนาคม 2569", 18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_page_break()

    # ════════════════════════════════════════
    # สารบัญ
    # ════════════════════════════════════════
    add_p(doc, "สารบัญ", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 6)

    toc_items = [
        "1. ปัญหาที่พบ — Deny Rate สูง สูญเสียรายได้",
        "2. Hospital Claim AI คืออะไร",
        "3. ฟีเจอร์หลัก 6 ด้าน",
        "4. ตัวอย่างการใช้งานจริง",
        "5. ความรู้ที่ AI ใช้ (Knowledge Base)",
        "6. เทคโนโลยีที่ใช้",
        "7. แผนการพัฒนา (Roadmap)",
        "8. ผลที่คาดหวัง (ROI)",
    ]
    for item in toc_items:
        add_p(doc, item, 16, space_after=8)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 1. ปัญหา
    # ════════════════════════════════════════
    add_p(doc, "1. ปัญหาที่พบ", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 6)

    add_p(doc, "สถานการณ์ปัจจุบัน:", 16, bold=True)
    add_bullet(doc, "Deny Rate เฉลี่ย 8-12% ของ claim ทั้งหมด")
    add_bullet(doc, "แผนก Cath Lab มี deny rate สูงกว่าค่าเฉลี่ย เนื่องจากเคสซับซ้อน มูลค่าสูง")
    add_bullet(doc, "สาเหตุ deny ที่พบบ่อย: เอกสารอุปกรณ์ไม่ครบ (HC09/HC13), ICD coding ไม่ตรง, ส่งเกินกำหนด")
    add_bullet(doc, "การอุทธรณ์ใช้เวลานาน ต้องค้นหาข้อมูลหลายระบบ")

    add_p(doc, "", 6)
    add_p(doc, "ผลกระทบ:", 16, bold=True)

    add_styled_table(doc,
        ["รายการ", "มูลค่าโดยประมาณ"],
        [
            ["เคส Cath Lab เฉลี่ย/เดือน", "30-50 เคส"],
            ["มูลค่าเฉลี่ย/เคส", "50,000-80,000 บาท"],
            ["Deny Rate ปัจจุบัน", "~10%"],
            ["รายได้ที่สูญเสีย/เดือน", "~150,000-400,000 บาท"],
            ["รายได้ที่สูญเสีย/ปี", "~1.8-4.8 ล้านบาท"],
        ],
        col_widths=[8, 8],
    )

    doc.add_page_break()

    # ════════════════════════════════════════
    # 2. Hospital Claim AI คืออะไร
    # ════════════════════════════════════════
    add_p(doc, "2. Hospital Claim AI คืออะไร", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 6)

    add_p(doc, "ระบบ AI ที่ช่วยเจ้าหน้าที่เบิกประกัน (Coder/UR) ทำงานได้เร็วขึ้น แม่นยำขึ้น โดย:", 14)
    add_p(doc, "", 4)

    add_number(doc, "ตรวจสอบ claim อัตโนมัติ ก่อนส่ง สปสช. — ลดโอกาสถูก deny", 14)
    add_number(doc, "วิเคราะห์สาเหตุ deny — หา root cause + แนะนำวิธีแก้", 14)
    add_number(doc, "ทำนายโอกาสถูก deny — ก่อนส่งเบิก รู้ล่วงหน้า", 14)
    add_number(doc, "ร่างหนังสืออุทธรณ์ — ภาษาราชการไทย พร้อมพิมพ์ลงนามส่ง", 14)
    add_number(doc, "แปลง Clinical Notes เป็น ICD codes — อัตโนมัติ", 14)
    add_number(doc, "ตรวจหลายเคสพร้อมกัน — เรียงตาม priority มูลค่า", 14)

    add_p(doc, "", 10)
    add_p(doc, "เป้าหมาย: ลด Deny Rate จาก 10% เหลือ 3%", 18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x00, 0x70, 0x30))

    doc.add_page_break()

    # ════════════════════════════════════════
    # 3. ฟีเจอร์หลัก
    # ════════════════════════════════════════
    add_p(doc, "3. ฟีเจอร์หลัก 6 ด้าน", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 6)

    # Feature 1
    add_p(doc, "3.1 ตรวจเคสก่อนส่งเบิก (Pre-submission Check)", 16, bold=True, color=RGBColor(0x44, 0x72, 0xC4))
    add_bullet(doc, "ตรวจ 8 Checkpoints อัตโนมัติ: ข้อมูลพื้นฐาน, Dx-Proc match, เอกสารอุปกรณ์, 16-File, Timing, CC/MCC, DRG, Drug/Lab")
    add_bullet(doc, "ให้คะแนน 0-100 พร้อมบอก CRITICAL / WARNING / PASS")
    add_bullet(doc, "แนะนำ CC/MCC optimization เพื่อเพิ่ม DRG weight")
    add_bullet(doc, "คำสั่ง: /claim check")
    add_p(doc, "", 6)

    # Feature 2
    add_p(doc, "3.2 วิเคราะห์เคสถูก Deny (Post-denial Analysis)", 16, bold=True, color=RGBColor(0x44, 0x72, 0xC4))
    add_bullet(doc, "รับ deny codes (HC09, HC13, IP01, C438, etc.) → หา root cause")
    add_bullet(doc, "จำแนก 6 categories: device, coding, drug, eligibility, timing, payment")
    add_bullet(doc, "แนะนำ fix steps ที่ actionable + ประเมิน recovery chance")
    add_bullet(doc, "คำสั่ง: /claim deny")
    add_p(doc, "", 6)

    # Feature 3
    add_p(doc, "3.3 ทำนาย Deny ก่อนส่ง (Deny Predictor)", 16, bold=True, color=RGBColor(0x44, 0x72, 0xC4))
    add_bullet(doc, "วิเคราะห์ 10 Risk Factors ให้คะแนน 0-100%")
    add_bullet(doc, "Verdict: SAFE (<20%), CAUTION (20-50%), HIGH_RISK (50-80%), ALMOST_CERTAIN (>80%)")
    add_bullet(doc, "บอก estimated loss if denied + recommendation")
    add_bullet(doc, "คำสั่ง: /claim predict")
    add_p(doc, "", 6)

    # Feature 4
    add_p(doc, "3.4 ร่างหนังสืออุทธรณ์ (Appeal Drafter)", 16, bold=True, color=RGBColor(0x44, 0x72, 0xC4))
    add_bullet(doc, "ร่างหนังสือภาษาราชการไทย พร้อมส่ง สปสช. เขต 6")
    add_bullet(doc, "อ้างอิงหลักเกณฑ์ สปสช. + เหตุผล clinical + guideline")
    add_bullet(doc, "แนะนำเอกสารแนบตามประเภท deny")
    add_bullet(doc, "สร้างไฟล์ DOCX พร้อมพิมพ์ลงนาม (ฟอนต์ TH SarabunPSK มาตรฐานราชการ)")
    add_bullet(doc, "คำสั่ง: /claim appeal")
    add_p(doc, "", 6)

    # Feature 5
    add_p(doc, "3.5 แปลง Clinical Notes เป็น ICD Codes (Smart Coder)", 16, bold=True, color=RGBColor(0x44, 0x72, 0xC4))
    add_bullet(doc, "รับ clinical notes ภาษาไทย/อังกฤษ → แนะนำ ICD-10-TM + ICD-9-CM")
    add_bullet(doc, "ระบุ PDx, SDx (พร้อม CC/MCC), Procedures")
    add_bullet(doc, "ประมาณ DRG + RW + expected payment")
    add_bullet(doc, "คำสั่ง: /claim code")
    add_p(doc, "", 6)

    # Feature 6
    add_p(doc, "3.6 ตรวจหลายเคสพร้อมกัน (Batch Optimizer)", 16, bold=True, color=RGBColor(0x44, 0x72, 0xC4))
    add_bullet(doc, "Upload CSV จาก e-Claim → AI ตรวจทุกเคส")
    add_bullet(doc, "จัดกลุ่ม: denied / at_risk / optimizable / ready")
    add_bullet(doc, "เรียง priority ตามมูลค่า (แก้เคสแพงก่อน)")
    add_bullet(doc, "สรุป total charge, denied, recoverable, optimization potential")
    add_bullet(doc, "คำสั่ง: /claim batch")

    doc.add_page_break()

    # ════════════════════════════════════════
    # 4. ตัวอย่างการใช้งาน
    # ════════════════════════════════════════
    add_p(doc, "4. ตัวอย่างการใช้งานจริง", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 6)

    add_p(doc, "เคส: AN 69-03556 — STEMI Anterior + Primary PCI", 16, bold=True)
    add_p(doc, "", 4)

    add_styled_table(doc,
        ["รายการ", "ข้อมูล"],
        [
            ["HN", "69-09349"],
            ["AN", "69-03556"],
            ["วันที่ Admit", "27 ก.พ. 2569"],
            ["วันที่ D/C", "1 มี.ค. 2569"],
            ["PDx", "I21.0 — Acute transmural MI of anterior wall"],
            ["Procedures", "36.06 (Stent insertion), 37.22 (L heart cath), 88.56 (Coronary arteriography)"],
            ["DRG", "05290 — Acute MI w single vessel PTCA wo sig CCC"],
            ["RW", "8.6544"],
            ["ยอดเรียกเก็บ", "71,322 บาท"],
        ],
        col_widths=[5, 11],
    )

    add_p(doc, "", 8)
    add_p(doc, "ผลจาก AI:", 16, bold=True)
    add_bullet(doc, "ตรวจ 8 checkpoints → พบ: HC09 (อุปกรณ์), HC13 (เอกสาร), Clopidogrel drug code ไม่ตรง")
    add_bullet(doc, "แนะนำแก้ ADP file (TYPE=5, Serial/Lot ตรง GPO VMI) + DRU file (TMT code)")
    add_bullet(doc, "Deny Prediction: HIGH_RISK 65% → หลังแก้: SAFE 12%")
    add_bullet(doc, "ร่างหนังสืออุทธรณ์ DOCX พร้อมพิมพ์ลงนาม")

    doc.add_page_break()

    # ════════════════════════════════════════
    # 5. Knowledge Base
    # ════════════════════════════════════════
    add_p(doc, "5. ความรู้ที่ AI ใช้ (Knowledge Base)", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 6)

    add_p(doc, "AI ไม่ได้เดา — ใช้ความรู้จากแหล่งจริงทั้งหมด:", 14)
    add_p(doc, "", 4)

    add_styled_table(doc,
        ["แหล่งความรู้", "จำนวน", "เนื้อหา"],
        [
            ["หลักเกณฑ์ สปสช.", "16 ไฟล์", "ประกาศ, ระเบียบ, drug catalog, instrument catalog"],
            ["ความรู้ตามแผนก", "11 ไฟล์", "Cath Lab, OR, Chemo, ICU, ER, Dialysis, OPD/NCD"],
            ["YouTube สปสช.", "5+ videos", "สกัด transcript จาก @nhsothailand"],
            ["Thai DRG v6.3.3", "Cardiac table", "RW, WtLOS, OT, payment estimate"],
            ["ICD-10-TM / ICD-9-CM", "Code sets", "Cardiac codes + CC/MCC mapping"],
            ["Deny Code Database", "15+ codes", "HC, IP, D, C codes + วิธีแก้"],
        ],
        col_widths=[5, 3, 8],
    )

    add_p(doc, "", 8)
    add_p(doc, "ครอบคลุม 9 แผนก:", 16, bold=True)

    add_styled_table(doc,
        ["แผนก", "Skill", "สถานะ"],
        [
            ["Cath Lab (สวนหัวใจ)", "cathlab-claim-checker", "พร้อมใช้"],
            ["OR/Surgery (ผ่าตัด)", "or-surgery-checker", "พร้อมใช้"],
            ["Chemo (เคมีบำบัด)", "chemo-checker", "พร้อมใช้"],
            ["ICU/NICU (วิกฤต)", "icu-checker", "พร้อมใช้"],
            ["ER/UCEP (ฉุกเฉิน)", "er-ucep-checker", "พร้อมใช้"],
            ["Dialysis (ฟอกไต)", "dialysis-checker", "พร้อมใช้"],
            ["OPD/NCD (โรคเรื้อรัง)", "opd-ncd-checker", "พร้อมใช้"],
            ["ทุกแผนก (ทั่วไป)", "claim-validator", "พร้อมใช้"],
            ["ICD Coding", "smart-coder + icd-coding", "พร้อมใช้"],
        ],
        col_widths=[5, 5, 3],
    )

    doc.add_page_break()

    # ════════════════════════════════════════
    # 6. เทคโนโลยี
    # ════════════════════════════════════════
    add_p(doc, "6. เทคโนโลยีที่ใช้", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 6)

    add_styled_table(doc,
        ["ส่วน", "เทคโนโลยี", "หน้าที่"],
        [
            ["AI Engine", "Claude API (Anthropic)", "วิเคราะห์เคส, สร้าง appeal, coding"],
            ["Rule Engine", "Python (8 checkpoints)", "ตรวจสอบ deterministic ไม่ต้องรอ AI"],
            ["Backend", "FastAPI + PostgreSQL", "API server + ฐานข้อมูล"],
            ["Frontend", "React + TailwindCSS", "Dashboard สำหรับเจ้าหน้าที่"],
            ["HIS Connector", "SSB REST API", "ดึงข้อมูลจาก HIS อัตโนมัติ"],
            ["Auth", "JWT (5 roles)", "admin, coder, dept_head, finance, readonly"],
            ["Notifications", "LINE Bot SDK", "แจ้งเตือนทีมเมื่อมี deny"],
            ["Reports", "Excel + DOCX", "รายงาน + หนังสืออุทธรณ์"],
        ],
        col_widths=[4, 5, 7],
    )

    add_p(doc, "", 10)
    add_p(doc, "สถาปัตยกรรมระบบ:", 16, bold=True)
    add_p(doc, "", 4)

    add_p(doc, "HIS (SSB)  →  AI Server  →  Dashboard (เจ้าหน้าที่)", 16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "                    ↓", 16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "           LINE Bot แจ้งเตือน", 16, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 7. Roadmap
    # ════════════════════════════════════════
    add_p(doc, "7. แผนการพัฒนา (Roadmap)", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 6)

    add_styled_table(doc,
        ["Phase", "ระยะเวลา", "สิ่งที่ทำ", "สถานะ"],
        [
            ["Phase 1", "เสร็จแล้ว", "Rule Engine 8 checkpoints + AI Engine + 20 Skills", "เสร็จแล้ว"],
            ["Phase 2", "เสร็จแล้ว", "API 7 endpoints + Deny Predictor + Smart Coder + Batch Optimizer", "เสร็จแล้ว"],
            ["Phase 3", "เสร็จแล้ว", "Appeal Drafter + DOCX generation + Knowledge Base 11 แผนก", "เสร็จแล้ว"],
            ["Phase 4", "กำลังทำ", "Dashboard (React) + LINE Bot notifications", "กำลังทำ"],
            ["Phase 5", "อนาคต", "เชื่อม HIS (SSB) intranet → auto check ทุกเคส", "วางแผน"],
            ["Phase 6", "อนาคต", "Full automation: auto check → auto fix → auto submit", "วางแผน"],
        ],
        col_widths=[3, 3, 7, 3],
    )

    add_p(doc, "", 10)
    add_p(doc, "แผน Automation 3 ระดับ:", 16, bold=True)
    add_p(doc, "", 4)

    add_styled_table(doc,
        ["ระดับ", "รูปแบบ", "ต้องทำ"],
        [
            ["Level 1: Semi-auto", "เจ้าหน้าที่พิมพ์คำสั่ง → AI วิเคราะห์", "ใช้ได้เลยตอนนี้"],
            ["Level 2: API auto", "HIS ส่งเคส → AI ตรวจ → แจ้ง LINE", "วาง AI server ใน LAN เดียวกับ SSB"],
            ["Level 3: Full auto", "Auto check → fix → submit → appeal", "ต้องมี human approval ก่อน submit"],
        ],
        col_widths=[4, 6, 6],
    )

    doc.add_page_break()

    # ════════════════════════════════════════
    # 8. ROI
    # ════════════════════════════════════════
    add_p(doc, "8. ผลที่คาดหวัง (ROI)", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 6)

    add_p(doc, "ถ้าลด Deny Rate จาก 10% เหลือ 3%:", 16, bold=True)
    add_p(doc, "", 4)

    add_styled_table(doc,
        ["ตัวชี้วัด", "ก่อนใช้ AI", "หลังใช้ AI", "ผลต่าง"],
        [
            ["Deny Rate", "~10%", "~3%", "-7%"],
            ["เคส deny/เดือน (50 เคส)", "~5 เคส", "~1.5 เคส", "-3.5 เคส"],
            ["มูลค่า deny/เดือน", "~300,000 บาท", "~90,000 บาท", "ลด 210,000 บาท"],
            ["รายได้เพิ่ม/ปี", "—", "—", "~2.5 ล้านบาท"],
            ["เวลาตรวจ claim/เคส", "15-30 นาที", "2-5 นาที", "ลด 80%"],
            ["เวลาร่างอุทธรณ์/เคส", "1-2 ชั่วโมง", "5 นาที", "ลด 90%"],
        ],
        col_widths=[5, 3.5, 3.5, 4],
        header_color="006030",
    )

    add_p(doc, "", 10)
    add_p(doc, "ประโยชน์เพิ่มเติม:", 16, bold=True)
    add_bullet(doc, "CC/MCC Optimization — เพิ่ม DRG weight ได้อีก 5-15% ของ RW ต่อเคส")
    add_bullet(doc, "ลดภาระเจ้าหน้าที่ — ลด burnout, ลดความผิดพลาดจากคน")
    add_bullet(doc, "Audit trail — ทุกเคสมี log ว่า AI แนะนำอะไร คนอนุมัติเมื่อไหร่")
    add_bullet(doc, "Knowledge สะสม — ยิ่งใช้ยิ่งฉลาด อัพเดตความรู้ได้ตลอด")

    add_p(doc, "", 20)
    add_p(doc, "รายได้เพิ่ม ~2.5 ล้านบาท/ปี  |  เวลาลด 80-90%", 20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x00, 0x70, 0x30))

    doc.add_page_break()

    # ════════════════════════════════════════
    # สรุป
    # ════════════════════════════════════════
    add_p(doc, "สรุป", 22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    add_p(doc, "", 10)

    add_p(doc, "Hospital Claim AI สำหรับ รพ.พญาไทศรีราชา", 18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "", 8)

    summary_items = [
        ("20 AI Skills", "ครอบคลุม 9 แผนก ตรวจ claim ทุกประเภท"),
        ("8 Checkpoints", "ตรวจอัตโนมัติก่อนส่งเบิก ลดโอกาส deny"),
        ("10 Risk Factors", "ทำนายโอกาสถูก deny ก่อนส่ง"),
        ("6 คำสั่ง /claim", "ใช้งานง่าย เจ้าหน้าที่เรียนรู้ได้เร็ว"),
        ("DOCX Appeal", "หนังสืออุทธรณ์พร้อมพิมพ์ลงนามส่ง สปสช."),
        ("Knowledge Base", "27+ ไฟล์ความรู้ อัพเดตได้ตลอด"),
        ("Semi-auto → Full auto", "เริ่มใช้ได้เลย พัฒนาต่อเป็น auto ได้"),
    ]

    for title, desc in summary_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run1 = p.add_run(f"  {title}  ")
        set_thai_font(run1, 16, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
        run2 = p.add_run(f"— {desc}")
        set_thai_font(run2, 16)

    add_p(doc, "", 20)
    add_p(doc, "ขอบคุณครับ", 24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1F, 0x4E, 0x79))

    # Save
    doc.save(output_path)
    print(f"สร้างเอกสาร presentation เสร็จ: {output_path}")
    print(f"ขนาดไฟล์: {os.path.getsize(output_path):,} bytes")


if __name__ == "__main__":
    import sys
    output = sys.argv[1] if len(sys.argv) > 1 else "docs/Hospital_Claim_AI_Presentation.docx"
    generate_presentation(output)
