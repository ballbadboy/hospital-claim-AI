"""Generate Appeal Letter as DOCX — เคส AN 69-03556"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def set_thai_font(run, size=14, bold=False):
    run.font.name = "TH SarabunPSK"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs", "TH SarabunPSK")


def add_thai_paragraph(doc, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_thai_font(run, size, bold)
    return p


def add_table_row(table, cells_data, bold_first=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        set_thai_font(run, 14, bold=(bold_first and i == 0))


def generate_appeal_docx(output_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Header
    add_thai_paragraph(doc, "ที่ พศ.11855/...........", 14, align=WD_ALIGN_PARAGRAPH.LEFT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("โรงพยาบาลพญาไทศรีราชา\n222 หมู่ 4 ถ.สุขุมวิท ต.ศรีราชา\nอ.ศรีราชา จ.ชลบุรี 20110")
    set_thai_font(run, 14)

    add_thai_paragraph(doc, "วันที่ .......... มีนาคม 2569", 14, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_thai_paragraph(doc, "", 8)

    # Subject
    add_thai_paragraph(doc, "เรื่อง  ขอทบทวนผลการตรวจสอบการขอรับค่าใช้จ่ายเพื่อบริการสาธารณสุข", 15, bold=True)
    add_thai_paragraph(doc, "เรียน  ผู้อำนวยการสำนักงานหลักประกันสุขภาพแห่งชาติ เขต 6 ระยอง", 15, bold=True)
    add_thai_paragraph(doc, "", 6)

    # Body - Intro
    add_thai_paragraph(doc,
        "        ตามที่โรงพยาบาลพญาไทศรีราชา รหัสหน่วยบริการ 11855 ได้ส่งข้อมูลการขอรับค่าใช้จ่าย"
        "เพื่อบริการสาธารณสุข สำหรับผู้ป่วยดังต่อไปนี้", 14)

    # Patient info table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    patient_data = [
        ("HN", "69-09349"),
        ("AN", "69-03556"),
        ("เลขประจำตัวประชาชน", "2411300021503"),
        ("วันที่รับเข้ารักษา", "27 กุมภาพันธ์ 2569 เวลา 11:44 น."),
        ("วันที่จำหน่าย", "1 มีนาคม 2569 เวลา 13:00 น."),
        ("การวินิจฉัยหลัก", "Acute transmural myocardial infarction of anterior wall (I21.0)"),
        ("หัตถการ", "Insertion of non-drug-eluting coronary artery stent (36.06)\n"
                    "Left heart catheterization (37.22)\n"
                    "Coronary arteriography using two catheters (88.56)"),
        ("DRG", "05290 (Acute MI w single vessel PTCA wo sig CCC)"),
        ("ค่าน้ำหนักสัมพัทธ์ (RW)", "8.6544"),
        ("ยอดเรียกเก็บ", "71,322.00 บาท"),
        ("สิทธิ", "หลักประกันสุขภาพแห่งชาติ (UCS)"),
    ]

    # Set header row
    hdr = table.rows[0].cells
    for i, text in enumerate(["รายการ", "ข้อมูล"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        set_thai_font(run, 14, bold=True)

    for label, value in patient_data:
        add_table_row(table, [label, value], bold_first=True)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    add_thai_paragraph(doc, "", 6)

    # Deny info
    add_thai_paragraph(doc,
        "        ทางโรงพยาบาลได้รับแจ้งว่าข้อมูลดังกล่าวไม่ผ่านการตรวจสอบ โดยมีรหัสปฏิเสธดังนี้", 14)

    # Deny table
    deny_table = doc.add_table(rows=1, cols=4)
    deny_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    deny_table.style = "Table Grid"

    deny_hdr = deny_table.rows[0].cells
    for i, text in enumerate(["รหัส", "ความหมาย", "รายการที่ถูกปฏิเสธ", "มูลค่า"]):
        p = deny_hdr[i].paragraphs[0]
        run = p.add_run(text)
        set_thai_font(run, 14, bold=True)

    deny_data = [
        ("HC09", "อุปกรณ์และอวัยวะเทียมในการทำหัตถการ", "INST (coronary stent)", "31,490 บาท"),
        ("HC13", "เงื่อนไข HC เพิ่มเติม", "เอกสารอุปกรณ์", "—"),
        ("IP01", "การจ่าย IP ในเขต", "IPINRGR", "—"),
    ]
    for row_data in deny_data:
        add_table_row(deny_table, row_data)

    add_thai_paragraph(doc, "", 6)
    add_thai_paragraph(doc,
        "        นอกจากนี้ รายการยา Clopidogrel ยังถูกปฏิเสธเนื่องจากรหัสยาไม่ตรง Drug Catalogue", 14)

    add_thai_paragraph(doc, "", 10)

    # Section: ข้อชี้แจง
    add_thai_paragraph(doc, "ข้อชี้แจง", 15, bold=True)

    add_thai_paragraph(doc, "1. ความจำเป็นทางคลินิก", 14, bold=True)
    add_thai_paragraph(doc,
        "        ผู้ป่วยรายนี้เข้ารับการรักษาด้วยอาการเจ็บหน้าอกรุนแรง ตรวจคลื่นไฟฟ้าหัวใจ (EKG) "
        "พบ ST segment elevation ใน lead V1-V4 ร่วมกับค่า Troponin สูงกว่าค่าปกติอย่างมีนัยสำคัญ "
        "วินิจฉัยเป็น ST-elevation myocardial infarction (STEMI) ของผนังหัวใจด้านหน้า "
        "จึงได้ดำเนินการทำ Primary Percutaneous Coronary Intervention (Primary PCI) "
        "ซึ่งเป็นมาตรฐานการรักษาตามแนวทางเวชปฏิบัติของสมาคมแพทย์โรคหัวใจแห่งประเทศไทย "
        "และ ACC/AHA 2025 Guideline for Acute Coronary Syndromes", 14)

    add_thai_paragraph(doc, "2. อุปกรณ์ที่ใช้ในการรักษา (HC09/HC13)", 14, bold=True)
    add_thai_paragraph(doc,
        "        ในการทำ PCI ได้ใช้ขดลวดค้ำยันหลอดเลือดโคโรนารี (Coronary stent) จำนวน 1 ชิ้น "
        "โดยอุปกรณ์ดังกล่าวได้จัดซื้อผ่านระบบ GPO VMI/SMI ตามระเบียบที่ สปสช. กำหนด "
        "ทางโรงพยาบาลได้ตรวจสอบและแก้ไขข้อมูลในแฟ้ม ADP ให้ถูกต้องครบถ้วนแล้ว ดังนี้", 14)

    bullet_items = [
        "TYPE = 5 (อวัยวะเทียม)",
        "CODE ตรงกับรายการที่ สปสช. กำหนด",
        "Serial/Lot number ตรงกับ GPO VMI record",
        "จำนวน 1 ชิ้น ตรงกับ Cath report",
    ]
    for item in bullet_items:
        p = doc.add_paragraph()
        p.style = "List Bullet"
        run = p.add_run(item)
        set_thai_font(run, 14)

    add_thai_paragraph(doc, "3. รายการยา Clopidogrel", 14, bold=True)
    add_thai_paragraph(doc,
        "        รายการยา Clopidogrel ที่ถูกปฏิเสธนั้น ทางโรงพยาบาลได้ตรวจสอบพบว่าเกิดจากรหัสยา"
        "ในระบบ HIS ไม่ตรงกับรหัส TMT ใน Drug Catalogue ของ สปสช. "
        "ซึ่งได้ดำเนินการแก้ไข mapping ระหว่าง HospDrugCode กับ TMTID ให้ถูกต้องเรียบร้อยแล้ว", 14)

    add_thai_paragraph(doc, "4. ข้อมูลที่แก้ไขแล้ว", 14, bold=True)
    add_thai_paragraph(doc,
        "        ทางโรงพยาบาลได้แก้ไขข้อมูลในแฟ้ม ADP (อุปกรณ์) และแฟ้ม DRU (ยา) "
        "ให้ถูกต้องตามมาตรฐานที่กำหนด และได้จัดส่งข้อมูลแก้ไขมาพร้อมหนังสือฉบับนี้", 14)

    add_thai_paragraph(doc, "", 10)

    # Attachments
    add_thai_paragraph(doc, "เอกสารแนบ", 15, bold=True)

    attachments = [
        "สำเนาเวชระเบียนผู้ป่วยใน (IPD record)",
        "รายงานผลการสวนหัวใจ (Cardiac Catheterization Report)",
        "ผลตรวจ EKG (ก่อนและหลังทำ PCI)",
        "ผลตรวจ Troponin",
        "ใบรายงานผลหัตถการ (Procedure Note)",
        "หลักฐานการจัดซื้อ stent ผ่าน GPO VMI (Lot/Serial confirmation)",
        "แฟ้ม ADP ที่แก้ไขแล้ว",
        "แฟ้ม DRU ที่แก้ไข Drug Catalogue mapping แล้ว",
        "ข้อมูล 16 แฟ้มที่แก้ไขแล้ว (.ecd file)",
    ]
    for i, item in enumerate(attachments, 1):
        add_thai_paragraph(doc, f"        {i}. {item}", 14)

    add_thai_paragraph(doc, "", 10)

    # Closing
    add_thai_paragraph(doc,
        "        จึงเรียนมาเพื่อโปรดพิจารณาทบทวนผลการตรวจสอบ และขอให้ดำเนินการจ่ายชดเชย"
        "ค่าบริการตามหลักเกณฑ์ที่กำหนด จะเป็นพระคุณยิ่ง", 14)

    add_thai_paragraph(doc, "", 20)
    add_thai_paragraph(doc, "ขอแสดงความนับถือ", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_thai_paragraph(doc, "", 30)

    add_thai_paragraph(doc, "(ลงนาม) .............................................", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_thai_paragraph(doc, "(................................................................)", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_thai_paragraph(doc, "ตำแหน่ง ผู้อำนวยการโรงพยาบาลพญาไทศรีราชา", 14, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    doc.save(output_path)
    print(f"✅ สร้างหนังสืออุทธรณ์เสร็จ: {output_path}")
    print(f"📄 ขนาดไฟล์: {os.path.getsize(output_path):,} bytes")


if __name__ == "__main__":
    import sys
    output = sys.argv[1] if len(sys.argv) > 1 else "docs/appeal_AN69-03556.docx"
    generate_appeal_docx(output)
