#!/usr/bin/env python3
"""
V993 Pre-Authorization Checker — ตรวจสอบ + อุทธรณ์เคส Cath Lab ที่ถูก deny จาก V993

Usage:
    python scripts/v993_checker.py <v993_deny.csv> [output_dir]

V993 = ระบบตรวจสอบ clinical indication ของ สปสช. สำหรับ Cath Lab
ต่างจาก e-Claim deny (HC09/IP01) ตรงที่ V993 ตรวจว่า "ทำหัตถการถูก indication ไหม"
"""

import csv, io, os, sys, time
from datetime import datetime, date
from pathlib import Path

try:
    import openpyxl
except ImportError:
    openpyxl = None

# ─── V993 Deny Patterns Knowledge Base ──────────────────

# Device-Procedure Matching (แจ้งจาก สปสช. มี.ค. 69)
DEVICE_PROC_MATCH = {
    "4307": "17.55", "4308": "17.55", "4309": "17.55",
    "4310": "17.55", "4320": "17.55", "4306": "36.06",
}

V993_DENY_RULES = {
    "(?:thromb|4310|17\\.55)": {
        "category": "DEVICE_PROC_MISMATCH",
        "severity": "CRITICAL",
        "meaning": "ใช้อุปกรณ์ Thrombectomy (4310) แต่ไม่มีรหัสหัตถการ 17.55 (Transluminal coronary atherectomy)",
        "root_cause": "OPR file ไม่มี ICD-9-CM 17.55 ซึ่ง สปสช. ทำ auto matching กับ ADP file",
        "appeal_strategy": "เพิ่มรหัส 17.55 ใน OPR file แล้วส่งข้อมูลใหม่ (ข้อมูลก่อน 17/3/69 ต้องส่งใหม่ทั้งหมด)",
        "appeal_success_rate": "สูงมาก (แก้ OPR file แล้วส่งใหม่)",
        "prevention": "ทุกเคสที่ใช้ device 4307-4310, 4320 ต้อง code 17.55 ใน OPR file เสมอ",
        "guideline_ref": "ICD-9-CM 2015: code 17.55 = Transluminal coronary atherectomy "
                         "(Rotational atherectomy, Cutting balloon, Thrombectomy) บังคับสำหรับ DRG v6",
        "appeal_points": [
            "อุปกรณ์ Thrombectomy Catheter (รหัส 4310) ที่ใช้ในหัตถการ "
            "เป็น Rheolytic thrombectomy ซึ่งจัดอยู่ในกลุ่ม Transluminal coronary atherectomy "
            "ตาม ICD-9-CM 2015",
            "ได้ดำเนินการเพิ่มรหัสหัตถการ 17.55 (Transluminal coronary atherectomy) "
            "ใน OPR file เรียบร้อยแล้ว พร้อมส่งข้อมูลใหม่",
            "ทั้งนี้การใช้ Thrombectomy device มีข้อบ่งชี้ทางคลินิกชัดเจน "
            "เพื่อกำจัดลิ่มเลือดในหลอดเลือดโคโรนารี ระหว่างทำ PCI "
            "ตาม ACC/AHA Guideline",
        ],
    },
    "Cx/OM-BASELINE: 51-75%": {
        "category": "STENOSIS_BORDERLINE",
        "severity": "HIGH",
        "meaning": "Stenosis 51-75% ไม่ถึงเกณฑ์ทำ PCI (ต้อง >70% หรือ FFR <0.80)",
        "root_cause": "ไม่มีผล FFR/iFR support ว่า lesion มี hemodynamic significance",
        "appeal_strategy": "ส่ง FFR/iFR report ที่แสดงค่า <0.80 หรือ iFR <0.89",
        "appeal_success_rate": "สูง (ถ้ามี FFR) / ต่ำ (ถ้าไม่มี)",
        "prevention": "ทำ FFR/iFR ทุกราย ที่ stenosis 50-70% ก่อนตัดสินใจ PCI",
        "guideline_ref": "ESC/EACTS 2024 Guideline: Class I recommendation for FFR/iFR in intermediate stenosis (50-70%)",
        "appeal_points": [
            "ผู้ป่วยรายนี้มี stenosis ที่ borderline (51-75%) ซึ่งตามแนวทางเวชปฏิบัติ ESC/EACTS 2024 "
            "แนะนำให้ประเมิน hemodynamic significance ด้วย FFR/iFR",
            "ผลการตรวจ FFR ได้ค่า [ระบุค่า FFR] ซึ่งต่ำกว่า 0.80 แสดงว่า lesion มี "
            "hemodynamic significance จึงมีข้อบ่งชี้ในการทำ PCI",
            "ทั้งนี้ visual assessment ของ stenosis อาจประเมินต่ำกว่าความเป็นจริง "
            "โดยเฉพาะใน eccentric lesion หรือ lesion ที่มี positive remodeling",
        ],
    },
    "NSTE-ACS.*multivessel": {
        "category": "MULTIVESSEL_SHOCK",
        "severity": "CRITICAL",
        "meaning": "ทำ multivessel PCI ของ non-culprit lesions พร้อมกัน ขณะมี cardiogenic shock",
        "root_cause": "ทำ complete revascularization ใน same setting ขณะ hemodynamically unstable",
        "appeal_strategy": "อ้าง clinical necessity: hemodynamic instability ต้องทำ complete revasc เพื่อ survival",
        "appeal_success_rate": "ปานกลาง (ต้องมี hemodynamic data สนับสนุน)",
        "prevention": "แยก staged PCI: ทำ culprit lesion ก่อน แล้ว non-culprit แยก admission",
        "guideline_ref": "CULPRIT-SHOCK Trial 2017: culprit-only PCI preferred in cardiogenic shock. "
                         "แต่ ACC/AHA 2021: complete revasc อาจ consider ใน selected cases",
        "appeal_points": [
            "ผู้ป่วยรายนี้มาด้วยภาวะ NSTE-ACS ร่วมกับ cardiogenic shock "
            "มี hemodynamic instability ที่รุนแรง (MAP [ระบุ], Lactate [ระบุ], vasopressor requirement)",
            "จากการประเมิน coronary angiography พบ multivessel disease ที่รุนแรง "
            "ทีม interventional cardiologist ประเมินแล้วว่าการทำ complete revascularization "
            "มีความจำเป็นเพื่อ hemodynamic support และ survival benefit",
            "อ้างอิง ACC/AHA 2021 Guideline ข้อ 7.3.2 ที่ระบุว่า complete revascularization "
            "อาจพิจารณาใน selected patients with cardiogenic shock เมื่อ operator มีประสบการณ์ "
            "และผู้ป่วยมี favorable anatomy",
            "ดังปรากฏตาม hemodynamic data และ cath report แนบ",
        ],
    },
    "SYNTAX-SCORE.*high.*>33": {
        "category": "HIGH_SYNTAX",
        "severity": "HIGH",
        "meaning": "SYNTAX score >33 ควร CABG ไม่ใช่ PCI ตาม guideline",
        "root_cause": "Complex coronary anatomy ที่ guideline แนะนำ CABG",
        "appeal_strategy": "อ้าง Heart Team decision + surgical risk สูง (STS score) + patient preference",
        "appeal_success_rate": "ปานกลาง (ต้องมี Heart Team note + STS score)",
        "prevention": "ทำ Heart Team discussion ทุกราย SYNTAX >22 + บันทึกเอกสาร",
        "guideline_ref": "ESC/EACTS 2024: SYNTAX >33 → CABG recommended (Class I). "
                         "แต่ PCI acceptable ถ้า surgical risk สูง หรือ patient refuses CABG",
        "appeal_points": [
            "ผู้ป่วยรายนี้มี SYNTAX score [ระบุค่า] ซึ่งจัดอยู่ในกลุ่ม high complexity "
            "ตามแนวทาง ESC/EACTS 2024 แนะนำ CABG เป็น preferred strategy",
            "อย่างไรก็ตาม ได้มีการประชุม Heart Team ประกอบด้วย interventional cardiologist "
            "cardiac surgeon และ clinical cardiologist เมื่อวันที่ [ระบุ] "
            "มีมติเลือก PCI เนื่องจาก [ระบุเหตุผล เช่น:]",
            "- ผู้ป่วยมี surgical risk สูง (STS score [ระบุ]%) จาก [ระบุ comorbidities]",
            "- ผู้ป่วยปฏิเสธการผ่าตัด CABG หลังได้รับคำอธิบายครบถ้วน (informed consent แนบ)",
            "- Coronary anatomy ที่แม้ SYNTAX score สูง แต่ target lesions สามารถทำ PCI ได้ "
            "ด้วย acceptable procedural risk",
            "ดังปรากฏตาม Heart Team meeting note และ cath report แนบ",
        ],
    },
}

# ─── Hospital Info ───────────────────────────────────────

HOSPITAL = {
    "hospcode": "11855",
    "hosp_name": "โรงพยาบาลพญาไทศรีราชา",
    "hosp_province": "จังหวัดชลบุรี",
    "nhso_region": "6",
    "nhso_region_name": "ระยอง",
}

# ─── CSV Parser ──────────────────────────────────────────

def _cell_to_str(val):
    """Convert cell value to string, handling dates and numbers."""
    if val is None:
        return ""
    if isinstance(val, (datetime, date)):
        return val.strftime("%d/%m/%Y")
    if isinstance(val, (int, float)):
        # PID ต้องเป็น string ไม่มีทศนิยม
        if val == int(val):
            return str(int(val))
        return str(val)
    return str(val).strip()


def _rows_to_cases(headers, data_rows):
    """Convert header + data rows → list of case dicts."""
    cases = []
    for row in data_rows:
        if len(row) < len(headers):
            row += [""] * (len(headers) - len(row))
        d = dict(zip(headers, row))
        if not d.get("PID"):
            continue
        cases.append({
            "hcode": d.get("HCODE", ""),
            "pid": d.get("PID", ""),
            "fname": d.get("FNAME", ""),
            "lname": d.get("LNAME", ""),
            "patient_name": f'{d.get("FNAME", "")} {d.get("LNAME", "")}'.strip(),
            "birthdate": d.get("BIRTHDATE", ""),
            "sex": d.get("SEX", ""),
            "admit": d.get("ADMISSION_DATE", ""),
            "case_type": d.get("TYPE", ""),
            "cag_status": d.get("CAG", ""),
            "cag_date": d.get("CAG_DATE", ""),
            "pci_status": d.get("PCI", ""),
            "pci_date": d.get("PCI_DATE", ""),
            "status": d.get("STATUS", ""),
            "doc_no": d.get("DOC_NO", ""),
            "deny_desc": d.get("DENY_DES", "").strip().strip('"').strip(),
        })
    return cases


def parse_v993_xlsx(file_path):
    """Parse V993 deny xlsx → list of case dicts (อ่าน xlsx ตรง)."""
    if openpyxl is None:
        raise ImportError("ต้อง pip install openpyxl ก่อน")
    wb = openpyxl.load_workbook(file_path, data_only=True)
    ws = wb.active

    # Read all rows as strings
    all_rows = []
    for row in ws.iter_rows(values_only=True):
        all_rows.append([_cell_to_str(c) for c in row])

    # Find header row
    header_idx = None
    for i, row in enumerate(all_rows):
        joined = " ".join(row)
        if "HCODE" in joined and "DENY_DES" in joined:
            header_idx = i
            break
    if header_idx is None:
        raise ValueError("ไม่พบ header row (HCODE/DENY_DES) ใน xlsx")

    headers = all_rows[header_idx]
    data_rows = all_rows[header_idx + 1:]
    return _rows_to_cases(headers, data_rows)


def parse_v993_csv(content):
    """Parse V993 deny CSV → list of case dicts."""
    reader = csv.reader(io.StringIO(content))
    rows = list(reader)

    # Find header
    header_idx = None
    for i, row in enumerate(rows):
        joined = ",".join(row)
        if "HCODE" in joined and "DENY_DES" in joined:
            header_idx = i
            break
    if header_idx is None:
        for i, row in enumerate(rows):
            if any("PID" in c for c in row):
                header_idx = i
                break
    if header_idx is None:
        raise ValueError("ไม่พบ header row (HCODE/PID/DENY_DES)")

    headers = [h.strip() for h in rows[header_idx]]
    data_rows = [[c.strip() for c in row] for row in rows[header_idx + 1:]]
    return _rows_to_cases(headers, data_rows)


def match_deny_pattern(deny_desc):
    """Match deny description to knowledge base pattern."""
    import re
    desc_lower = deny_desc.lower()
    for pattern, rule in V993_DENY_RULES.items():
        if re.search(pattern.lower(), desc_lower):
            return rule
    # Fallback
    return {
        "category": "OTHER",
        "severity": "MEDIUM",
        "meaning": deny_desc,
        "root_cause": "ต้องตรวจสอบรายละเอียดเพิ่มเติม",
        "appeal_strategy": "อ้าง clinical justification พร้อมเอกสาร",
        "appeal_success_rate": "ไม่แน่นอน",
        "prevention": "ตรวจสอบ indication ก่อนทำหัตถการ",
        "guideline_ref": "ESC/EACTS Guidelines",
        "appeal_points": [f"ข้อชี้แจงสำหรับ: {deny_desc}"],
    }


def _fill_appeal_points(appeal_points, cag_appeal):
    """Replace [ระบุ...] placeholders in appeal points with CAG form data."""
    if not cag_appeal:
        return appeal_points

    replacements = {
        "[ระบุค่า FFR]": cag_appeal.get("ffr_value", "[ระบุค่า FFR]"),
        "[ระบุค่า]": cag_appeal.get("syntax_score", "[ระบุค่า]"),
        "[ระบุ comorbidities]": cag_appeal.get("comorbidities", "[ระบุ comorbidities]"),
    }

    # MAP/Lactate for shock cases
    if cag_appeal.get("shock"):
        replacements["MAP [ระบุ]"] = f'MAP {cag_appeal.get("map_value", "[ระบุ]")}'
        replacements["Lactate [ระบุ]"] = f'Lactate {cag_appeal.get("lactate_value", "[ระบุ]")}'

    # STS score
    replacements["STS score [ระบุ]%"] = f'STS score {cag_appeal.get("sts_score", "[ระบุ]")}%'

    # SYNTAX score
    if cag_appeal.get("syntax_score"):
        replacements["SYNTAX score [ระบุค่า]"] = f'SYNTAX score {cag_appeal["syntax_score"]}'

    # Physician / date
    if cag_appeal.get("physician"):
        replacements["เมื่อวันที่ [ระบุ]"] = f'เมื่อวันที่ {cag_appeal.get("cag_date", "[ระบุ]")} โดย {cag_appeal["physician"]}'

    filled = []
    for point in appeal_points:
        for placeholder, value in replacements.items():
            point = point.replace(placeholder, str(value))
        filled.append(point)
    return filled


def analyze_v993_case(case, cag_data=None):
    """Analyze single V993 case → findings + appeal strategy."""
    rule = match_deny_pattern(case["deny_desc"])

    # Calculate age — handle both ค.ศ. and พ.ศ.
    age = ""
    try:
        bd = case["birthdate"]
        if "/" in bd:
            parts = bd.split("/")
            birth_year = int(parts[2])
            # พ.ศ. > 2100 แปลงเป็น ค.ศ. (ปี ค.ศ. ปัจจุบัน ~2026 ไม่มีทางเกิน 2100)
            if birth_year > 2100:
                birth_year -= 543
            age = datetime.now().year - birth_year
    except Exception:
        pass

    # Fill appeal points from CAG form data if available
    filled_rule = dict(rule)
    if cag_data:
        from scripts.cag_form_parser import cag_to_appeal_data
        cag_appeal = cag_to_appeal_data(cag_data)
        filled_rule["appeal_points"] = _fill_appeal_points(rule["appeal_points"], cag_appeal)
        # Add CAG data to case for report
        case["cag_form_data"] = cag_appeal

    return {
        **case,
        "age": age,
        "rule": filled_rule,
        "category": filled_rule["category"],
        "severity": filled_rule["severity"],
    }


# ─── DOCX Generation ────────────────────────────────────

def generate_v993_report(cases_analyzed, output_path):
    """Generate V993 summary report DOCX."""
    project_root = Path(__file__).resolve().parent.parent
    sys.path.insert(0, str(project_root))
    from scripts.docx_generators import _setup_doc, _p, _p2, _section_header, _kv_table, _data_table, _bullet_list, _number_list
    from scripts.docx_generators import RED, DARK_RED, ORANGE, GREEN, NAVY, DARK_BLUE, WHITE, GRAY, BLACK
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _setup_doc("CONFIDENTIAL - V993 Pre-Auth Deny Analysis")
    today = datetime.now().strftime("%d/%m/%Y")

    # Title
    for _ in range(2):
        doc.add_paragraph()
    _p(doc, 'V993 Pre-Authorization Deny Analysis', 28, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, after=4)
    _p(doc, f'{HOSPITAL["hosp_name"]} ({HOSPITAL["hospcode"]})', 18, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, after=4)
    _p(doc, f'{len(cases_analyzed)} cases | {today}', 16, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, after=20)
    doc.add_page_break()

    # Summary table
    _section_header(doc, '1. Summary')
    _p(doc, f'Total cases: {len(cases_analyzed)}', 16, True, after=4)

    # Category breakdown
    cats = {}
    for c in cases_analyzed:
        cat = c["category"]
        cats[cat] = cats.get(cat, 0) + 1
    cat_labels = {
        "STENOSIS_BORDERLINE": "Stenosis 51-75% (ไม่ถึงเกณฑ์)",
        "MULTIVESSEL_SHOCK": "Multivessel PCI ใน shock",
        "HIGH_SYNTAX": "SYNTAX >33 (ควร CABG)",
        "OTHER": "อื่นๆ",
    }
    _data_table(doc,
        ["Deny Pattern", "จำนวน", "ระดับ", "โอกาสอุทธรณ์สำเร็จ"],
        [[cat_labels.get(k, k), str(v),
          "CRITICAL" if k == "MULTIVESSEL_SHOCK" else "HIGH",
          {"STENOSIS_BORDERLINE": "สูง (ถ้ามี FFR)", "MULTIVESSEL_SHOCK": "ปานกลาง", "HIGH_SYNTAX": "ปานกลาง"}.get(k, "-")]
         for k, v in cats.items()],
        col_widths=[Cm(5), Cm(2), Cm(3), Cm(5.5)])

    # Prevention recommendations
    _section_header(doc, '2. Systemic Prevention')
    _number_list(doc, [
        "FFR Protocol: stenosis 50-70% ต้องทำ FFR/iFR ก่อน PCI ทุกราย",
        "Heart Team: SYNTAX >22 ต้องผ่าน Heart Team discussion + บันทึก",
        "Staged PCI: ห้ามทำ multivessel non-culprit PCI ใน same setting",
        "Pre-Auth Checklist: ตรวจ stenosis >70% / FFR <0.80 / SYNTAX ≤33 / Heart Team note",
    ], 14)

    # Each case detail
    _section_header(doc, '3. Case Details')
    for i, c in enumerate(cases_analyzed, 1):
        r = c["rule"]
        _p(doc, f'Case {i}: {c["patient_name"]} ({c["sex"]}, {c["age"]} ปี)',
           16, True, DARK_BLUE, before=12, after=4)

        _kv_table(doc, [
            ("PID", c["pid"]),
            ("Admit", c["admit"]),
            ("Type", c["case_type"]),
            ("CAG", c["cag_status"]),
            ("PCI", c["pci_status"]),
            ("Deny", c["deny_desc"]),
            ("Category", f'{c["category"]} ({r["severity"]})'),
        ], Cm(4), Cm(12))

        _p(doc, 'Root Cause:', 14, True, DARK_RED, before=8, after=2)
        _p(doc, r["root_cause"], 14, after=4)

        _p(doc, 'Appeal Strategy:', 14, True, GREEN, before=4, after=2)
        _p(doc, r["appeal_strategy"], 14, after=4)

        _p(doc, 'Guideline Reference:', 14, True, NAVY, before=4, after=2)
        _p(doc, r["guideline_ref"], 13, after=4)

        _p(doc, 'Prevention:', 14, True, ORANGE, before=4, after=2)
        _p(doc, r["prevention"], 14, after=8)

    # Signature
    doc.add_paragraph()
    _p(doc, f'Report generated: {today} | Hospital Claim AI', 12, False, GRAY, WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(output_path)
    return output_path


def generate_v993_appeal(case_analyzed, output_path):
    """Generate V993 appeal letter DOCX for single case."""
    project_root = Path(__file__).resolve().parent.parent
    sys.path.insert(0, str(project_root))
    from scripts.docx_generators import _setup_doc, _p, _p2, _section_header, _kv_table, _bullet_list, _number_list
    from scripts.docx_generators import RED, NAVY, GREEN, GRAY, BLACK, ORANGE
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    c = case_analyzed
    r = c["rule"]
    today = datetime.now().strftime("%d/%m/%Y")
    doc = _setup_doc("CONFIDENTIAL")
    sec = doc.sections[0]
    sec.left_margin = Cm(3)

    # Header
    _p(doc, f'ที่ รพ. {HOSPITAL["hospcode"]}/......', 16, after=0)
    _p(doc, f'{HOSPITAL["hosp_name"]}\n{HOSPITAL["hosp_province"]}', 16, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    _p(doc, f'วันที่ ...... ................. 2569', 16, align=WD_ALIGN_PARAGRAPH.RIGHT, after=12)

    # Subject
    _p2(doc, [
        ('เรื่อง  ', 16, True, BLACK),
        ('ขอทบทวนผลการตรวจสอบการให้บริการหัตถการสวนหัวใจและหลอดเลือด (V993)', 16, False, BLACK),
    ], after=4)

    _p2(doc, [
        ('เรียน  ', 16, True, BLACK),
        (f'ผู้อำนวยการสำนักงานหลักประกันสุขภาพแห่งชาติ เขต {HOSPITAL["nhso_region"]} {HOSPITAL["nhso_region_name"]}', 16, False, BLACK),
    ], after=12)

    # Opening
    _p(doc,
       f'ตามที่{HOSPITAL["hosp_name"]} รหัสหน่วยบริการ {HOSPITAL["hospcode"]} '
       'ได้ให้บริการหัตถการสวนหัวใจและหลอดเลือดแก่ผู้ป่วยรายต่อไปนี้ '
       'และได้รับแจ้งว่าไม่ผ่านการตรวจสอบ Pre-Authorization (V993) '
       'ทั้งนี้ ขอชี้แจงดังนี้',
       16, after=12, indent=2.5)

    # Patient info
    _p(doc, 'ข้อมูลผู้ป่วย', 16, True, after=4, before=4)
    _kv_table(doc, [
        ('ชื่อ-สกุล', c["patient_name"]),
        ('เลขประจำตัวประชาชน', c["pid"]),
        ('เพศ / อายุ', f'{c["sex"]} / {c["age"]} ปี'),
        ('วันที่รับเข้า', c["admit"]),
        ('ประเภท', c["case_type"]),
        ('CAG', f'{c["cag_status"]} ({c["cag_date"]})'),
        ('PCI', f'{c["pci_status"]} ({c["pci_date"]})'),
        ('เหตุผลที่ไม่ผ่าน', c["deny_desc"]),
    ], Cm(5), Cm(10.5))

    # Justification
    doc.add_paragraph()
    _p(doc, 'ข้อชี้แจง', 16, True, after=8, before=12)

    for i, point in enumerate(r["appeal_points"], 1):
        _p(doc, f'{i}. {point}', 16, after=6, indent=2.5)

    # Guideline reference
    _p(doc, 'อ้างอิงแนวทางเวชปฏิบัติ', 16, True, after=4, before=12)
    _p(doc, r["guideline_ref"], 16, after=8, indent=2.5)

    # Attachments
    _p(doc, 'เอกสารแนบ:', 16, True, after=4, before=12)

    attachments = [
        "สำเนาเวชระเบียนผู้ป่วย",
        "รายงานผลการสวนหัวใจ (Cardiac Catheterization Report)",
        "ภาพ Angiogram ที่แสดง lesion",
    ]
    if c["category"] == "STENOSIS_BORDERLINE":
        attachments.extend(["ผลการตรวจ FFR/iFR", "ภาพ Pressure wire recording"])
    elif c["category"] == "MULTIVESSEL_SHOCK":
        attachments.extend(["Hemodynamic data (MAP, Lactate, Vasopressor)", "ICU chart"])
    elif c["category"] == "HIGH_SYNTAX":
        attachments.extend(["Heart Team meeting note", "STS Risk Score calculation", "Informed consent (ถ้าผู้ป่วยปฏิเสธ CABG)"])

    _number_list(doc, attachments, 16)

    # Closing
    _p(doc,
       'จึงเรียนมาเพื่อโปรดพิจารณาทบทวนผลการตรวจสอบ '
       'และขอให้ดำเนินการจ่ายชดเชยค่าบริการตามหลักเกณฑ์ที่กำหนด '
       'จะเป็นพระคุณยิ่ง',
       16, after=24, before=12, indent=2.5)

    _p(doc, 'ขอแสดงความนับถือ', 16, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=36)
    _p(doc, '(............................................................)', 16, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    _p(doc, f'ตำแหน่ง ผู้อำนวยการ{HOSPITAL["hosp_name"]}', 16, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)

    doc.save(output_path)
    return output_path


# ─── Main Pipeline ───────────────────────────────────────

def _load_cag_forms(cag_dir):
    """Load all CAG form DOCX files from a directory, keyed by PID."""
    if not cag_dir or not Path(cag_dir).is_dir():
        return {}
    from scripts.cag_form_parser import parse_cag_form
    forms = {}
    for f in Path(cag_dir).glob("*.docx"):
        if f.name.startswith("~$"):
            continue
        try:
            data = parse_cag_form(str(f))
            pid = data["general"].get("pid", "")
            if pid:
                forms[pid] = data
                print(f"  CAG form loaded: {f.name} (PID: {pid})")
        except Exception as e:
            print(f"  CAG form skip: {f.name} ({e})")
    return forms


def run_v993_pipeline(csv_path, output_dir=None, cag_dir=None):
    """Main: CSV/XLSX → parse → (optional) match CAG forms → analyze → generate DOCX."""
    t0 = time.time()
    file_ext = Path(csv_path).suffix.lower()

    if file_ext in (".xlsx", ".xls"):
        cases = parse_v993_xlsx(csv_path)
    else:
        with open(csv_path, encoding="utf-8-sig") as f:
            content = f.read()
        cases = parse_v993_csv(content)

    # Load CAG forms if directory provided
    cag_forms = _load_cag_forms(cag_dir)
    if cag_forms:
        print(f"  CAG forms matched: {len(cag_forms)} files")

    analyzed = [analyze_v993_case(c, cag_forms.get(c["pid"])) for c in cases]
    t_parse = time.time() - t0

    if not output_dir:
        output_dir = Path(csv_path).parent / "v993_reports"
    output_dir = Path(output_dir)
    output_dir.mkdir(exist_ok=True)

    # Generate summary report
    t1 = time.time()
    report_path = output_dir / "V993_Summary_Report.docx"
    generate_v993_report(analyzed, str(report_path))

    # Generate individual appeal letters
    for c in analyzed:
        pid_safe = c["pid"][-4:]
        appeal_path = output_dir / f'V993_Appeal_{pid_safe}_{c["category"]}.docx'
        generate_v993_appeal(c, str(appeal_path))

    t_gen = time.time() - t1

    print(f"\n{'='*60}")
    print(f"  V993 PRE-AUTH DENY — PIPELINE RESULT")
    print(f"{'='*60}")
    print(f"  CSV        : {csv_path}")
    print(f"  Cases      : {len(analyzed)}")
    print(f"  Parse time : {t_parse:.3f}s")
    print(f"  DOCX time  : {t_gen:.3f}s ({1 + len(analyzed)} files)")
    print(f"  Total time : {t_parse + t_gen:.3f}s")
    print(f"  Output     : {output_dir}")
    print(f"{'─'*60}")
    print(f"  CASES:")
    for i, c in enumerate(analyzed, 1):
        print(f"  {i}. {c['patient_name']:<20} {c['category']:<25} {c['rule']['severity']}")
    print(f"{'='*60}")

    return analyzed


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/v993_checker.py <v993_deny.csv/xlsx> [output_dir] [--cag <cag_forms_dir>]")
        print("  --cag <dir>  โฟลเดอร์ที่มี CAG Data Form DOCX ที่กรอกแล้ว (match by PID)")
        sys.exit(1)

    csv_path = sys.argv[1]
    out_dir = None
    cag_dir = None

    # Parse args
    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == "--cag" and i + 1 < len(sys.argv):
            cag_dir = sys.argv[i + 1]
            i += 2
        else:
            out_dir = sys.argv[i]
            i += 1

    run_v993_pipeline(csv_path, out_dir, cag_dir)
