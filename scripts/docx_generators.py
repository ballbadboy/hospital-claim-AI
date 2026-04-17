#!/usr/bin/env python3
"""
Reusable DOCX generators for Hospital Claim AI.
Usage:
    from docx_generators import generate_deny_report, generate_appeal_letter
    generate_deny_report(data, "/path/to/output.docx")
    generate_appeal_letter(data, "/path/to/output.docx")

Or CLI:
    python docx_generators.py report  '{"an":"[REDACTED-AN]",...}' output.docx
    python docx_generators.py appeal  '{"an":"[REDACTED-AN]",...}' output.docx
"""

import json, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Theme ───────────────────────────────────────────
RED      = RGBColor(0xDC, 0x26, 0x26)
DARK_RED = RGBColor(0xB9, 0x1C, 0x1C)
ORANGE   = RGBColor(0xF5, 0x7C, 0x00)
GREEN    = RGBColor(0x2E, 0x7D, 0x32)
NAVY     = RGBColor(0x0D, 0x47, 0xA1)
DARK_BLUE= RGBColor(0x1A, 0x23, 0x7E)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
GRAY     = RGBColor(0x61, 0x61, 0x61)
BLACK    = RGBColor(0x21, 0x21, 0x21)
FONT     = 'TH Sarabun New'

# ─── Low-level helpers ───────────────────────────────
def _f(run, size=14, bold=False, color=BLACK):
    """Apply Thai font to a run."""
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    return run

def _shd(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def _border(cell, bottom="E0E0E0"):
    tc = cell._tc.get_or_add_tcPr()
    tc.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{bottom}"/>'
        f'<w:top w:val="nil" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'<w:left w:val="nil" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'<w:right w:val="nil" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'</w:tcBorders>'))

def _p(doc, text, size=14, bold=False, color=BLACK,
       align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, indent=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    _f(p.add_run(text), size, bold, color)
    return p

def _p2(doc, parts, before=0, after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add paragraph with multiple styled runs. parts = [(text, size, bold, color), ...]"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    for text, size, bold, color in parts:
        _f(p.add_run(text), size, bold, color)
    return p

def _section_header(doc, text):
    """Navy section header with bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="0D47A1"/>'
        f'</w:pBdr>'))
    _f(p.add_run(text), 18, True, NAVY)

def _kv_table(doc, rows, lw=Cm(5), rw=Cm(11)):
    """Key-value table with light borders."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(rows):
        cl, cv = t.rows[i].cells[0], t.rows[i].cells[1]
        _f(cl.paragraphs[0].add_run(k), 14, True, GRAY)
        _f(cv.paragraphs[0].add_run(str(v)), 14, False, BLACK)
        cl.width, cv.width = lw, rw
        for c in (cl, cv):
            _border(c)
    return t

def _data_table(doc, headers, rows, header_bg="1A237E", row_bgs=None, col_widths=None):
    """Generic table with colored header."""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        _shd(cell, header_bg)
        r = cell.paragraphs[0].add_run(h)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _f(r, 13, True, WHITE)
    # Data
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = t.rows[i+1].cells[j]
            if row_bgs and i < len(row_bgs):
                _shd(cell, row_bgs[i])
            _f(cell.paragraphs[0].add_run(str(val)), 12)
    # Widths
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w
    return t

def _bullet_list(doc, items, size=13):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        _f(p.add_run(item), size)

def _number_list(doc, items, size=13):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        _f(p.add_run(item), size)

def _setup_doc(confidential_text="CONFIDENTIAL \u2014 Hospital Claim AI System"):
    """Create doc with A4 page, Thai font default, header/footer."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2)
    sec.left_margin, sec.right_margin = Cm(2.5), Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # Header
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(confidential_text)
    r.font.name = FONT; r.font.size = Pt(10); r.font.color.rgb = GRAY; r.italic = True

    # Footer with page number
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run('Hospital Claim AI \u2014 ')
    r.font.name = FONT; r.font.size = Pt(10); r.font.color.rgb = GRAY
    for ftype in ('begin', None, 'end'):
        rr = fp.add_run()
        if ftype:
            rr._r.append(parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="{ftype}"/>'))
        else:
            rr._r.append(parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    return doc


# ═══════════════════════════════════════════════════════
#  DENY REPORT GENERATOR
# ═══════════════════════════════════════════════════════

def generate_deny_report(d: dict, output_path: str) -> str:
    """
    Generate Cath Lab Deny Report DOCX.

    Required keys in d:
        rep_no, hn, an, pid, right, admit, discharge, los,
        hcode, hmain, href, drg, rw, status,
        central_reimburse, charge, expected_payment, actual, loss,
        issues: [{ id, title, fund, subfund, deny_code, causes:[], fixes:[] }],
        actions: [{ priority, item, action, impact }],
        recommendation_a, recommendation_b,
        report_date
    """
    doc = _setup_doc()

    # ── Title Page ──
    for _ in range(3):
        doc.add_paragraph()
    _p(doc, '\u2588'*30, 8, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    _p(doc, '\u0e23\u0e32\u0e22\u0e07\u0e32\u0e19\u0e01\u0e32\u0e23\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a',
       28, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _p(doc, 'Cath Lab Claim \u2014 Deny Analysis',
       24, True, DARK_BLUE, WD_ALIGN_PARAGRAPH.CENTER, after=4)
    _p(doc, '\u2588'*30, 8, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    _p(doc, f'AN: {d["an"]}', 20, True, BLACK, WD_ALIGN_PARAGRAPH.CENTER, after=4)
    _p(doc, f'HN: {d["hn"]}', 16, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, after=4)
    _p(doc, f'\u26d4  DENY \u2014 \u0e2a\u0e39\u0e0d\u0e40\u0e2a\u0e35\u0e22 {d["loss"]} \u0e1a\u0e32\u0e17',
       22, True, RED, WD_ALIGN_PARAGRAPH.CENTER, before=12, after=12)
    _p(doc, f'\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48\u0e23\u0e32\u0e22\u0e07\u0e32\u0e19: {d["report_date"]}',
       14, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _p(doc, '\u0e2a\u0e23\u0e49\u0e32\u0e07\u0e42\u0e14\u0e22: Hospital Claim AI System',
       14, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, after=2)
    doc.add_page_break()

    # ── 1. Case Summary ──
    _section_header(doc, '1. \u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e40\u0e04\u0e2a (Case Summary)')
    _kv_table(doc, [
        ('REP No.', d.get('rep_no','-')), ('HN', d['hn']), ('AN', d['an']),
        ('PID', d.get('pid','-')), ('\u0e2a\u0e34\u0e17\u0e18\u0e34\u0e4c', d.get('right','-')),
        ('\u0e27\u0e31\u0e19\u0e40\u0e02\u0e49\u0e32\u0e23\u0e31\u0e01\u0e29\u0e32', d['admit']),
        ('\u0e27\u0e31\u0e19\u0e08\u0e33\u0e2b\u0e19\u0e48\u0e32\u0e22', d['discharge']),
        ('LOS', f'{d["los"]} \u0e27\u0e31\u0e19'),
        ('HCODE', d.get('hcode','-')), ('HMAIN', d.get('hmain','-')),
        ('HREF', d.get('href','-')), ('DRG', d['drg']),
        ('RW', str(d['rw'])), ('Status', d.get('status','DENY')),
    ])

    # ── 2. Financial Impact ──
    _section_header(doc, '2. \u0e1c\u0e25\u0e01\u0e23\u0e30\u0e17\u0e1a\u0e17\u0e32\u0e07\u0e01\u0e32\u0e23\u0e40\u0e07\u0e34\u0e19 (Financial Impact)')
    fin_rows = [
        ('Central Reimburse', f'{d["central_reimburse"]} \u0e1a\u0e32\u0e17'),
        ('\u0e04\u0e48\u0e32\u0e23\u0e31\u0e01\u0e29\u0e32', f'{d["charge"]} \u0e1a\u0e32\u0e17'),
        ('Expected Payment', f'{d["expected_payment"]} \u0e1a\u0e32\u0e17'),
        ('Actual Received', f'{d["actual"]} \u0e1a\u0e32\u0e17'),
        ('\u0e2a\u0e39\u0e0d\u0e40\u0e2a\u0e35\u0e22 (LOSS)', f'{d["loss"]} \u0e1a\u0e32\u0e17'),
    ]
    ft = doc.add_table(rows=len(fin_rows), cols=2)
    for i, (k, v) in enumerate(fin_rows):
        cl, cv = ft.rows[i].cells[0], ft.rows[i].cells[1]
        _f(cl.paragraphs[0].add_run(k), 14, True)
        cv.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = cv.paragraphs[0].add_run(v)
        _f(r, 14)
        cl.width, cv.width = Cm(7), Cm(9)
        if i == len(fin_rows)-1:  # LOSS row
            _shd(cl, 'FFEBEE'); _shd(cv, 'FFEBEE')
            for pp in (cl.paragraphs[0], cv.paragraphs[0]):
                for rr in pp.runs:
                    rr.font.color.rgb = DARK_RED; rr.bold = True
        for c in (cl, cv):
            _border(c)

    # ── 3. Deny Analysis ──
    _section_header(doc, '3. \u0e27\u0e34\u0e40\u0e04\u0e23\u0e32\u0e30\u0e2b\u0e4c\u0e2a\u0e32\u0e40\u0e2b\u0e15\u0e38 Deny')
    for issue in d.get('issues', []):
        severity_color = RED if issue.get('severity','critical') == 'critical' else ORANGE
        _p2(doc, [
            (f'\u0e1b\u0e31\u0e0d\u0e2b\u0e32\u0e17\u0e35\u0e48 {issue["id"]}: ', 16, True, severity_color),
            (issue['title'], 16, True, BLACK),
        ], before=12, after=4)

        it = doc.add_table(rows=3, cols=2)
        for j, (k, v) in enumerate([
            ('\u0e01\u0e2d\u0e07\u0e17\u0e38\u0e19', issue.get('fund','-')),
            ('Sub-fund', issue.get('subfund','-')),
            ('Deny Code', issue.get('deny_code','-')),
        ]):
            _f(it.rows[j].cells[0].paragraphs[0].add_run(k), 13, True)
            _f(it.rows[j].cells[1].paragraphs[0].add_run(v), 13)
            it.rows[j].cells[0].width = Cm(4)
            for c in (it.rows[j].cells[0], it.rows[j].cells[1]):
                _border(c)

        _p(doc, '\u0e2a\u0e32\u0e40\u0e2b\u0e15\u0e38\u0e17\u0e35\u0e48\u0e40\u0e1b\u0e47\u0e19\u0e44\u0e1b\u0e44\u0e14\u0e49:',
           14, True, DARK_RED, before=8)
        _bullet_list(doc, issue.get('causes', []))

        _p(doc, '\u0e27\u0e34\u0e18\u0e35\u0e41\u0e01\u0e49\u0e44\u0e02:',
           14, True, GREEN, before=8)
        _number_list(doc, issue.get('fixes', []))

    # ── 4. Action Plan ──
    doc.add_page_break()
    _section_header(doc, '4. \u0e41\u0e1c\u0e19\u0e01\u0e32\u0e23\u0e41\u0e01\u0e49\u0e44\u0e02 (Action Plan)')
    actions = d.get('actions', [])
    bg_cycle = ['FFEBEE', 'FFF3E0', 'FFF8E1', 'E8F5E9']
    _data_table(doc,
        ['\u0e25\u0e33\u0e14\u0e31\u0e1a', '\u0e23\u0e32\u0e22\u0e01\u0e32\u0e23',
         '\u0e2a\u0e34\u0e48\u0e07\u0e17\u0e35\u0e48\u0e15\u0e49\u0e2d\u0e07\u0e17\u0e33', 'Impact'],
        [[a['priority'], a['item'], a['action'], a['impact']] for a in actions],
        row_bgs=[bg_cycle[i % len(bg_cycle)] for i in range(len(actions))],
        col_widths=[Cm(3), Cm(4), Cm(5.5), Cm(3)])

    # ── 5. Recommendation ──
    _section_header(doc, '5. \u0e02\u0e49\u0e2d\u0e40\u0e2a\u0e19\u0e2d\u0e41\u0e19\u0e30')
    # Box
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="8" w:color="1A237E"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="8" w:color="1A237E"/>'
        f'<w:left w:val="single" w:sz="12" w:space="8" w:color="1A237E"/>'
        f'<w:right w:val="single" w:sz="12" w:space="8" w:color="1A237E"/>'
        f'</w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EAF6"/>'))
    _f(p.add_run(f'\u0e21\u0e39\u0e25\u0e04\u0e48\u0e32\u0e40\u0e04\u0e2a: ~{d["loss"]} \u0e1a\u0e32\u0e17 \u2014 \u0e04\u0e38\u0e49\u0e21\u0e04\u0e48\u0e32\u0e21\u0e32\u0e01\u0e17\u0e35\u0e48\u0e08\u0e30\u0e41\u0e01\u0e49\u0e44\u0e02'),
       16, True, NAVY)

    if d.get('recommendation_a'):
        _p2(doc, [
            ('\u0e17\u0e32\u0e07\u0e40\u0e25\u0e37\u0e2d\u0e01 A (\u0e41\u0e19\u0e30\u0e19\u0e33): ', 14, True, GREEN),
            (d['recommendation_a'], 14, False, BLACK),
        ], before=8)
    if d.get('recommendation_b'):
        _p2(doc, [
            ('\u0e17\u0e32\u0e07\u0e40\u0e25\u0e37\u0e2d\u0e01 B: ', 14, True, ORANGE),
            (d['recommendation_b'], 14, False, BLACK),
        ], before=4)

    # ── Signature ──
    doc.add_paragraph()
    _p(doc, f'\u0e23\u0e32\u0e22\u0e07\u0e32\u0e19\u0e2a\u0e23\u0e49\u0e32\u0e07\u0e42\u0e14\u0e22: Hospital Claim AI System | {d["report_date"]}',
       12, False, GRAY, WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════
#  APPEAL LETTER GENERATOR
# ═══════════════════════════════════════════════════════

def generate_appeal_letter(d: dict, output_path: str) -> str:
    """
    Generate formal appeal letter DOCX.

    Required keys in d:
        hospcode, hosp_name, hosp_province, nhso_region, nhso_region_name,
        patient_name, pid, hn, an, right,
        admit, discharge, pdx, pdx_name, procedures, drg, rw,
        charge, deny_codes: [{ code, meaning, item, value }],
        justifications: [{ title, text }],
        corrections: [],
        attachments: [],
        report_date
    """
    doc = _setup_doc("CONFIDENTIAL \u2014 Hospital Claim AI")
    sec = doc.sections[0]
    sec.left_margin = Cm(3)  # Gov standard

    # ── Header ──
    _p(doc, f'\u0e17\u0e35\u0e48 \u0e23\u0e1e. {d["hospcode"]}/\u2026\u2026\u2026\u2026', 16, after=0)
    _p(doc, f'{d["hosp_name"]}\n{d["hosp_province"]}',
       16, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    _p(doc, f'\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48 {d["report_date"]}',
       16, align=WD_ALIGN_PARAGRAPH.RIGHT, after=12)

    # ── Subject ──
    _p2(doc, [
        ('\u0e40\u0e23\u0e37\u0e48\u0e2d\u0e07  ', 16, True, BLACK),
        ('\u0e02\u0e2d\u0e17\u0e1a\u0e17\u0e27\u0e19\u0e1c\u0e25\u0e01\u0e32\u0e23\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a\u0e01\u0e32\u0e23\u0e02\u0e2d\u0e23\u0e31\u0e1a\u0e04\u0e48\u0e32\u0e43\u0e0a\u0e49\u0e08\u0e48\u0e32\u0e22\u0e40\u0e1e\u0e37\u0e48\u0e2d\u0e1a\u0e23\u0e34\u0e01\u0e32\u0e23\u0e2a\u0e32\u0e18\u0e32\u0e23\u0e13\u0e2a\u0e38\u0e02', 16, False, BLACK),
    ], after=4)

    _p2(doc, [
        ('\u0e40\u0e23\u0e35\u0e22\u0e19  ', 16, True, BLACK),
        (f'\u0e1c\u0e39\u0e49\u0e2d\u0e33\u0e19\u0e27\u0e22\u0e01\u0e32\u0e23\u0e2a\u0e33\u0e19\u0e31\u0e01\u0e07\u0e32\u0e19\u0e2b\u0e25\u0e31\u0e01\u0e1b\u0e23\u0e30\u0e01\u0e31\u0e19\u0e2a\u0e38\u0e02\u0e20\u0e32\u0e1e\u0e41\u0e2b\u0e48\u0e07\u0e0a\u0e32\u0e15\u0e34 \u0e40\u0e02\u0e15 {d["nhso_region"]} {d["nhso_region_name"]}', 16, False, BLACK),
    ], after=12)

    # ── Opening ──
    _p(doc,
       f'\u0e15\u0e32\u0e21\u0e17\u0e35\u0e48{d["hosp_name"]} \u0e23\u0e2b\u0e31\u0e2a\u0e2b\u0e19\u0e48\u0e27\u0e22\u0e1a\u0e23\u0e34\u0e01\u0e32\u0e23 {d["hospcode"]} '
       '\u0e44\u0e14\u0e49\u0e2a\u0e48\u0e07\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e01\u0e32\u0e23\u0e02\u0e2d\u0e23\u0e31\u0e1a\u0e04\u0e48\u0e32\u0e43\u0e0a\u0e49\u0e08\u0e48\u0e32\u0e22\u0e40\u0e1e\u0e37\u0e48\u0e2d\u0e1a\u0e23\u0e34\u0e01\u0e32\u0e23\u0e2a\u0e32\u0e18\u0e32\u0e23\u0e13\u0e2a\u0e38\u0e02 '
       '\u0e2a\u0e33\u0e2b\u0e23\u0e31\u0e1a\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e23\u0e32\u0e22\u0e15\u0e48\u0e2d\u0e44\u0e1b\u0e19\u0e35\u0e49 '
       '\u0e1c\u0e48\u0e32\u0e19\u0e23\u0e30\u0e1a\u0e1a e-Claim \u0e41\u0e25\u0e49\u0e27 '
       '\u0e17\u0e32\u0e07\u0e42\u0e23\u0e07\u0e1e\u0e22\u0e32\u0e1a\u0e32\u0e25\u0e44\u0e14\u0e49\u0e23\u0e31\u0e1a\u0e41\u0e08\u0e49\u0e07\u0e27\u0e48\u0e32\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e14\u0e31\u0e07\u0e01\u0e25\u0e48\u0e32\u0e27'
       '\u0e44\u0e21\u0e48\u0e1c\u0e48\u0e32\u0e19\u0e01\u0e32\u0e23\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a '
       '\u0e17\u0e31\u0e49\u0e07\u0e19\u0e35\u0e49 \u0e02\u0e2d\u0e0a\u0e35\u0e49\u0e41\u0e08\u0e07\u0e14\u0e31\u0e07\u0e19\u0e35\u0e49',
       16, after=12, indent=2.5)

    # ── Patient Table ──
    _p(doc, '\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22', 16, True, after=4, before=4)
    pt_rows = [
        ('\u0e0a\u0e37\u0e48\u0e2d-\u0e2a\u0e01\u0e38\u0e25', d.get('patient_name','[.............]')),
        ('CID', d.get('pid','-')), ('HN', d['hn']), ('AN', d['an']),
        ('\u0e2a\u0e34\u0e17\u0e18\u0e34\u0e4c', d.get('right','-')),
        ('\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48\u0e23\u0e31\u0e1a\u0e40\u0e02\u0e49\u0e32', d['admit']),
        ('\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48\u0e08\u0e33\u0e2b\u0e19\u0e48\u0e32\u0e22', d['discharge']),
        ('PDx', f'{d["pdx"]} {d.get("pdx_name","")}'),
        ('\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23', d.get('procedures','-')),
        ('DRG', d['drg']), ('RW', str(d['rw'])),
        ('\u0e22\u0e2d\u0e14\u0e40\u0e1a\u0e34\u0e01', f'{d["charge"]} \u0e1a\u0e32\u0e17'),
    ]
    pt = doc.add_table(rows=len(pt_rows), cols=2)
    pt.style = 'Table Grid'
    for i, (k, v) in enumerate(pt_rows):
        _f(pt.rows[i].cells[0].paragraphs[0].add_run(k), 14, True)
        _f(pt.rows[i].cells[1].paragraphs[0].add_run(str(v)), 14)
        pt.rows[i].cells[0].width = Cm(5)
        _shd(pt.rows[i].cells[0], 'F5F5F5')

    # ── Deny Table ──
    doc.add_paragraph()
    _p(doc, '\u0e23\u0e32\u0e22\u0e25\u0e30\u0e40\u0e2d\u0e35\u0e22\u0e14\u0e01\u0e32\u0e23\u0e1b\u0e0f\u0e34\u0e40\u0e2a\u0e18\u0e01\u0e32\u0e23\u0e08\u0e48\u0e32\u0e22\u0e0a\u0e14\u0e40\u0e0a\u0e22',
       16, True, after=4, before=8)
    deny_codes = d.get('deny_codes', [])
    if deny_codes:
        _data_table(doc,
            ['\u0e23\u0e2b\u0e31\u0e2a', '\u0e04\u0e27\u0e32\u0e21\u0e2b\u0e21\u0e32\u0e22',
             '\u0e23\u0e32\u0e22\u0e01\u0e32\u0e23', '\u0e21\u0e39\u0e25\u0e04\u0e48\u0e32'],
            [[dc['code'], dc['meaning'], dc['item'], dc['value']] for dc in deny_codes],
            col_widths=[Cm(2.5), Cm(4), Cm(5), Cm(4)])

    # ── Justifications ──
    doc.add_paragraph()
    _p(doc, '\u0e17\u0e31\u0e49\u0e07\u0e19\u0e35\u0e49 \u0e02\u0e2d\u0e0a\u0e35\u0e49\u0e41\u0e08\u0e07\u0e14\u0e31\u0e07\u0e19\u0e35\u0e49',
       16, True, after=8, before=12)
    for j in d.get('justifications', []):
        _p(doc, j['title'], 16, True, after=4, before=8)
        _p(doc, j['text'], 16, after=8, indent=2.5)

    # ── Corrections ──
    if d.get('corrections'):
        _p(doc, '\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e17\u0e35\u0e48\u0e44\u0e14\u0e49\u0e14\u0e33\u0e40\u0e19\u0e34\u0e19\u0e01\u0e32\u0e23\u0e41\u0e01\u0e49\u0e44\u0e02',
           16, True, after=4, before=8)
        _bullet_list(doc, d['corrections'], 16)

    # ── Closing ──
    _p(doc,
       '\u0e08\u0e36\u0e07\u0e40\u0e23\u0e35\u0e22\u0e19\u0e21\u0e32\u0e40\u0e1e\u0e37\u0e48\u0e2d\u0e42\u0e1b\u0e23\u0e14\u0e1e\u0e34\u0e08\u0e32\u0e23\u0e13\u0e32\u0e17\u0e1a\u0e17\u0e27\u0e19\u0e1c\u0e25\u0e01\u0e32\u0e23\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a '
       '\u0e41\u0e25\u0e30\u0e02\u0e2d\u0e43\u0e2b\u0e49\u0e14\u0e33\u0e40\u0e19\u0e34\u0e19\u0e01\u0e32\u0e23\u0e08\u0e48\u0e32\u0e22\u0e0a\u0e14\u0e40\u0e0a\u0e22\u0e04\u0e48\u0e32\u0e1a\u0e23\u0e34\u0e01\u0e32\u0e23\u0e15\u0e32\u0e21\u0e2b\u0e25\u0e31\u0e01\u0e40\u0e01\u0e13\u0e11\u0e4c\u0e17\u0e35\u0e48\u0e01\u0e33\u0e2b\u0e19\u0e14 '
       '\u0e08\u0e30\u0e40\u0e1b\u0e47\u0e19\u0e1e\u0e23\u0e30\u0e04\u0e38\u0e13\u0e22\u0e34\u0e48\u0e07',
       16, after=24, before=12, indent=2.5)

    _p(doc, '\u0e02\u0e2d\u0e41\u0e2a\u0e14\u0e07\u0e04\u0e27\u0e32\u0e21\u0e19\u0e31\u0e1a\u0e16\u0e37\u0e2d',
       16, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=36)
    _p(doc, '(\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026)',
       16, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    _p(doc, f'\u0e15\u0e33\u0e41\u0e2b\u0e19\u0e48\u0e07 \u0e1c\u0e39\u0e49\u0e2d\u0e33\u0e19\u0e27\u0e22\u0e01\u0e32\u0e23{d["hosp_name"]}',
       16, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)

    # ── Attachments ──
    if d.get('attachments'):
        doc.add_paragraph()
        _p(doc, '\u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23\u0e41\u0e19\u0e1a:', 16, True, before=12, after=4)
        for i, att in enumerate(d['attachments'], 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(2)
            _f(p.add_run(f'{i}. {att}'), 16)

    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════
#  CLI
# ═══════════════════════════════════════════════════════

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: python docx_generators.py <report|appeal> '<json>' output.docx")
        sys.exit(1)

    cmd = sys.argv[1]
    data = json.loads(sys.argv[2])
    out = sys.argv[3]

    if cmd == 'report':
        generate_deny_report(data, out)
    elif cmd == 'appeal':
        generate_appeal_letter(data, out)
    else:
        print(f"Unknown command: {cmd}")
        sys.exit(1)

    print(f"Generated: {out}")
