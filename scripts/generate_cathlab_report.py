#!/usr/bin/env python3
"""Generate Cath Lab Deny Analysis Report as DOCX"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# === Colors ===
RED = RGBColor(0xDC, 0x26, 0x26)
DARK_RED = RGBColor(0xB9, 0x1C, 0x1C)
ORANGE = RGBColor(0xF5, 0x7C, 0x00)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
DARK_BLUE = RGBColor(0x1A, 0x23, 0x7E)
NAVY = RGBColor(0x0D, 0x47, 0xA1)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x61, 0x61, 0x61)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
BLACK = RGBColor(0x21, 0x21, 0x21)

# === Helpers ===
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val["sz"]}" '
            f'w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_paragraph(doc, text, font_name='TH Sarabun New', font_size=14,
                         bold=False, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_colored_run(paragraph, text, font_name='TH Sarabun New', font_size=14,
                    bold=False, color=BLACK):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run

def add_section_header(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="{str(color).replace("#", "")}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    return p

def create_info_table(doc, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(rows_data):
        # Label cell
        cell_l = table.rows[i].cells[0]
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = GRAY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        if col_widths:
            cell_l.width = col_widths[0]
        # Value cell
        cell_v = table.rows[i].cells[1]
        p = cell_v.paragraphs[0]
        run = p.add_run(str(value))
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(14)
        run.font.color.rgb = BLACK
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        if col_widths:
            cell_v.width = col_widths[1]
        # Light border bottom
        for cell in [cell_l, cell_v]:
            set_cell_border(cell,
                bottom={"val": "single", "sz": "4", "color": "E0E0E0"},
                top={"val": "nil", "sz": "0", "color": "FFFFFF"},
                left={"val": "nil", "sz": "0", "color": "FFFFFF"},
                right={"val": "nil", "sz": "0", "color": "FFFFFF"})
    return table

# === MAIN ===
def generate_report():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- Default Style ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)
    font.color.rgb = BLACK
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # --- Header ---
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run('CONFIDENTIAL \u2014 Hospital Claim AI System')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.italic = True

    # --- Footer ---
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('Hospital Claim AI \u2014 Cath Lab Deny Analysis Report \u2014 ')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    # Page number
    run2 = footer_para.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run2._r.append(fldChar1)
    run3 = footer_para.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run3._r.append(instrText)
    run4 = footer_para.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run4._r.append(fldChar2)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    doc.add_paragraph()  # spacer
    doc.add_paragraph()  # spacer
    doc.add_paragraph()  # spacer

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('\u2588' * 30)
    run.font.color.rgb = NAVY
    run.font.size = Pt(8)

    add_styled_paragraph(doc, '\u0e23\u0e32\u0e22\u0e07\u0e32\u0e19\u0e01\u0e32\u0e23\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a',
                         font_size=28, bold=True, color=NAVY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_styled_paragraph(doc, 'Cath Lab Claim \u2014 Deny Analysis',
                         font_size=24, bold=True, color=DARK_BLUE,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run('\u2588' * 30)
    run.font.color.rgb = NAVY
    run.font.size = Pt(8)

    # Key info on title page
    add_styled_paragraph(doc, 'AN: [REDACTED-AN]', font_size=20, bold=True, color=BLACK,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, 'HN: [REDACTED-HN]', font_size=16, color=GRAY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # Status badge
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('\u26d4  DENY \u2014 \u0e2a\u0e39\u0e0d\u0e40\u0e2a\u0e35\u0e22 72,264 \u0e1a\u0e32\u0e17')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RED

    add_styled_paragraph(doc, '\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48\u0e23\u0e32\u0e22\u0e07\u0e32\u0e19: 31 \u0e21\u0e35\u0e19\u0e32\u0e04\u0e21 2569',
                         font_size=14, color=GRAY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_styled_paragraph(doc, '\u0e2a\u0e23\u0e49\u0e32\u0e07\u0e42\u0e14\u0e22: Hospital Claim AI System',
                         font_size=14, color=GRAY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    # Page break
    doc.add_page_break()

    # ============================================================
    # SECTION 1: CASE SUMMARY
    # ============================================================
    add_section_header(doc, '1. \u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e40\u0e04\u0e2a (Case Summary)')

    info_rows = [
        ('REP No.', '690300013'),
        ('HN', '[REDACTED-HN]'),
        ('AN', '[REDACTED-AN]'),
        ('PID', '[REDACTED-CID]'),
        ('\u0e2a\u0e34\u0e17\u0e18\u0e34\u0e4c', 'UC (UCS)'),
        ('\u0e27\u0e31\u0e19\u0e40\u0e02\u0e49\u0e32\u0e23\u0e31\u0e01\u0e29\u0e32 (Admit)', '27/02/2569 \u0e40\u0e27\u0e25\u0e32 11:44'),
        ('\u0e27\u0e31\u0e19\u0e08\u0e33\u0e2b\u0e19\u0e48\u0e32\u0e22 (Discharge)', '01/03/2569 \u0e40\u0e27\u0e25\u0e32 13:00'),
        ('LOS', '2 \u0e27\u0e31\u0e19'),
        ('HCODE (\u0e23\u0e1e.\u0e17\u0e35\u0e48\u0e23\u0e31\u0e01\u0e29\u0e32)', '10825'),
        ('HMAIN (\u0e23\u0e1e.\u0e15\u0e49\u0e19\u0e2a\u0e31\u0e07\u0e01\u0e31\u0e14)', '11855'),
        ('HREF', '89 (Refer case)'),
        ('DRG', '05290'),
        ('RW', '8.6544'),
        ('Status', 'DENY'),
    ]
    create_info_table(doc, info_rows, col_widths=[Cm(5), Cm(11)])

    # ============================================================
    # SECTION 2: FINANCIAL IMPACT
    # ============================================================
    add_section_header(doc, '2. \u0e1c\u0e25\u0e01\u0e23\u0e30\u0e17\u0e1a\u0e17\u0e32\u0e07\u0e01\u0e32\u0e23\u0e40\u0e07\u0e34\u0e19 (Financial Impact)')

    # Financial table with colored cells
    fin_table = doc.add_table(rows=6, cols=2)
    fin_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fin_data = [
        ('\u0e22\u0e2d\u0e14 Central Reimburse', '71,322 \u0e1a\u0e32\u0e17'),
        ('\u0e04\u0e48\u0e32\u0e23\u0e31\u0e01\u0e29\u0e32\u0e17\u0e35\u0e48\u0e40\u0e23\u0e35\u0e22\u0e01\u0e40\u0e01\u0e47\u0e1a', '31,517 \u0e1a\u0e32\u0e17'),
        ('Expected Payment (RW\u00d7Base)', '8.6544 \u00d7 8,350 = 72,264 \u0e1a\u0e32\u0e17'),
        ('Actual Received', '0 \u0e1a\u0e32\u0e17'),
        ('Base Rate (\u0e43\u0e19\u0e40\u0e02\u0e15)', '8,350 \u0e1a\u0e32\u0e17/AdjRW'),
        ('\u0e2a\u0e39\u0e0d\u0e40\u0e2a\u0e35\u0e22 (LOSS)', '72,264 \u0e1a\u0e32\u0e17'),
    ]

    for i, (label, value) in enumerate(fin_data):
        cell_l = fin_table.rows[i].cells[0]
        cell_v = fin_table.rows[i].cells[1]

        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(14)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        cell_l.width = Cm(7)

        p = cell_v.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(value)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        cell_v.width = Cm(9)

        # Last row = LOSS → red background
        if i == 5:
            set_cell_shading(cell_l, 'FFEBEE')
            set_cell_shading(cell_v, 'FFEBEE')
            for p2 in [cell_l.paragraphs[0], cell_v.paragraphs[0]]:
                for r in p2.runs:
                    r.font.color.rgb = DARK_RED
                    r.bold = True
        # Row 3 = Actual = 0
        if i == 3:
            for p2 in [cell_v.paragraphs[0]]:
                for r in p2.runs:
                    r.font.color.rgb = RED
                    r.bold = True

        for cell in [cell_l, cell_v]:
            set_cell_border(cell,
                bottom={"val": "single", "sz": "4", "color": "E0E0E0"},
                top={"val": "nil", "sz": "0", "color": "FFFFFF"},
                left={"val": "nil", "sz": "0", "color": "FFFFFF"},
                right={"val": "nil", "sz": "0", "color": "FFFFFF"})

    # ============================================================
    # SECTION 3: DENY ANALYSIS
    # ============================================================
    add_section_header(doc, '3. \u0e27\u0e34\u0e40\u0e04\u0e23\u0e32\u0e30\u0e2b\u0e4c\u0e2a\u0e32\u0e40\u0e2b\u0e15\u0e38 Deny (Root Cause Analysis)')

    # --- Issue 1 ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('\u0e1b\u0e31\u0e0d\u0e2b\u0e32\u0e17\u0e35\u0e48 1: ')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RED
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    run2 = p.add_run('HC09 \u2014 \u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c/INST (Deny Code: I04)')
    run2.font.name = 'TH Sarabun New'
    run2.font.size = Pt(16)
    run2.bold = True
    run2.font.color.rgb = BLACK
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    issue1_table = doc.add_table(rows=3, cols=2)
    issue1_data = [
        ('\u0e01\u0e2d\u0e07\u0e17\u0e38\u0e19', 'HC09 (\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e41\u0e25\u0e30\u0e2d\u0e27\u0e31\u0e22\u0e27\u0e30\u0e40\u0e17\u0e35\u0e22\u0e21\u0e43\u0e19\u0e01\u0e32\u0e23\u0e17\u0e33\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23)'),
        ('Sub-fund', 'INST'),
        ('Deny Code', 'I04'),
    ]
    for i, (label, value) in enumerate(issue1_data):
        cell_l = issue1_table.rows[i].cells[0]
        cell_v = issue1_table.rows[i].cells[1]
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        cell_l.width = Cm(4)
        p = cell_v.paragraphs[0]
        run = p.add_run(value)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        for cell in [cell_l, cell_v]:
            set_cell_border(cell,
                bottom={"val": "single", "sz": "2", "color": "E0E0E0"},
                top={"val": "nil", "sz": "0", "color": "FFFFFF"},
                left={"val": "nil", "sz": "0", "color": "FFFFFF"},
                right={"val": "nil", "sz": "0", "color": "FFFFFF"})

    add_styled_paragraph(doc, '\u0e2a\u0e32\u0e40\u0e2b\u0e15\u0e38\u0e17\u0e35\u0e48\u0e40\u0e1b\u0e47\u0e19\u0e44\u0e1b\u0e44\u0e14\u0e49:',
                         font_size=14, bold=True, color=DARK_RED, space_before=8)
    causes_1 = [
        '\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25 ADP file \u0e44\u0e21\u0e48\u0e04\u0e23\u0e1a/\u0e44\u0e21\u0e48\u0e16\u0e39\u0e01\u0e15\u0e49\u0e2d\u0e07 (TYPE, CODE, Serial/Lot Number)',
        '\u0e23\u0e32\u0e04\u0e32\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e44\u0e21\u0e48\u0e15\u0e23\u0e07\u0e01\u0e31\u0e1a Fee Schedule \u0e17\u0e35\u0e48\u0e01\u0e33\u0e2b\u0e19\u0e14',
        'Device \u0e44\u0e21\u0e48\u0e2d\u0e22\u0e39\u0e48\u0e43\u0e19\u0e1a\u0e31\u0e0d\u0e0a\u0e35\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e17\u0e35\u0e48 \u0e2a\u0e1b\u0e2a\u0e0a. \u0e2d\u0e19\u0e38\u0e21\u0e31\u0e15\u0e34',
    ]
    for c in causes_1:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(13)
        # Clear and re-add with font
        p.clear()
        run = p.add_run(c)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    add_styled_paragraph(doc, '\u0e27\u0e34\u0e18\u0e35\u0e41\u0e01\u0e49\u0e44\u0e02:',
                         font_size=14, bold=True, color=GREEN, space_before=8)
    fixes_1 = [
        '\u0e40\u0e1b\u0e34\u0e14 ADP file \u0e15\u0e23\u0e27\u0e08 field: TYPE, CODE, QTY, RATE, SERIALNO',
        '\u0e15\u0e23\u0e27\u0e08\u0e27\u0e48\u0e32 stent code \u0e15\u0e23\u0e07\u0e01\u0e31\u0e1a NHSO Device Catalog',
        '\u0e40\u0e1e\u0e34\u0e48\u0e21 lot number / serial number \u0e16\u0e49\u0e32\u0e02\u0e32\u0e14',
        '\u0e15\u0e23\u0e27\u0e08\u0e23\u0e32\u0e04\u0e32\u0e27\u0e48\u0e32\u0e44\u0e21\u0e48\u0e40\u0e01\u0e34\u0e19\u0e40\u0e1e\u0e14\u0e32\u0e19\u0e17\u0e35\u0e48\u0e01\u0e33\u0e2b\u0e19\u0e14',
        '\u0e2a\u0e48\u0e07\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e43\u0e2b\u0e21\u0e48 (resubmit) \u0e2b\u0e23\u0e37\u0e2d\u0e2d\u0e38\u0e17\u0e18\u0e23\u0e13\u0e4c\u0e1e\u0e23\u0e49\u0e2d\u0e21\u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23',
    ]
    for idx, f in enumerate(fixes_1, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        run = p.add_run(f)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # --- Issue 2 ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('\u0e1b\u0e31\u0e0d\u0e2b\u0e32\u0e17\u0e35\u0e48 2: ')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RED
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    run2 = p.add_run('IP01 \u2014 \u0e40\u0e01\u0e13\u0e11\u0e4c\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e43\u0e19 (Deny Code: C)')
    run2.font.name = 'TH Sarabun New'
    run2.font.size = Pt(16)
    run2.bold = True
    run2.font.color.rgb = BLACK
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    issue2_table = doc.add_table(rows=3, cols=2)
    issue2_data = [
        ('\u0e01\u0e2d\u0e07\u0e17\u0e38\u0e19', 'IP01 (IP \u0e40\u0e01\u0e13\u0e11\u0e4c)'),
        ('Sub-fund', 'IPINRGR (IP \u0e43\u0e19\u0e40\u0e02\u0e15)'),
        ('Deny Code', 'C (\u0e44\u0e21\u0e48\u0e1c\u0e48\u0e32\u0e19\u0e01\u0e32\u0e23\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a)'),
    ]
    for i, (label, value) in enumerate(issue2_data):
        cell_l = issue2_table.rows[i].cells[0]
        cell_v = issue2_table.rows[i].cells[1]
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        cell_l.width = Cm(4)
        p = cell_v.paragraphs[0]
        run = p.add_run(value)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        for cell in [cell_l, cell_v]:
            set_cell_border(cell,
                bottom={"val": "single", "sz": "2", "color": "E0E0E0"},
                top={"val": "nil", "sz": "0", "color": "FFFFFF"},
                left={"val": "nil", "sz": "0", "color": "FFFFFF"},
                right={"val": "nil", "sz": "0", "color": "FFFFFF"})

    add_styled_paragraph(doc, '\u0e2a\u0e32\u0e40\u0e2b\u0e15\u0e38\u0e17\u0e35\u0e48\u0e40\u0e1b\u0e47\u0e19\u0e44\u0e1b\u0e44\u0e14\u0e49:',
                         font_size=14, bold=True, color=DARK_RED, space_before=8)
    causes_2 = [
        'LOS 2 \u0e27\u0e31\u0e19 \u0e2d\u0e32\u0e08\u0e2a\u0e31\u0e49\u0e19\u0e40\u0e01\u0e34\u0e19\u0e44\u0e1b\u0e2a\u0e33\u0e2b\u0e23\u0e31\u0e1a DRG \u0e19\u0e35\u0e49',
        '\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25 IPD file \u0e44\u0e21\u0e48\u0e04\u0e23\u0e1a (DATEADM, DATEDSC, DISCHT)',
        'Authen Code \u0e2b\u0e21\u0e14\u0e2d\u0e32\u0e22\u0e38\u0e2b\u0e23\u0e37\u0e2d\u0e44\u0e21\u0e48\u0e15\u0e23\u0e07\u0e01\u0e31\u0e1a\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48 admit',
        '\u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23 Refer \u0e44\u0e21\u0e48\u0e04\u0e23\u0e1a (HMAIN 11855 \u2260 HCODE 10825 \u2192 \u0e40\u0e04\u0e2a Refer)',
        'HMAIN2 = 06 \u0e1c\u0e34\u0e14\u0e1b\u0e01\u0e15\u0e34 (\u0e1b\u0e01\u0e15\u0e34\u0e15\u0e49\u0e2d\u0e07 5 \u0e2b\u0e25\u0e31\u0e01) \u2014 \u0e2d\u0e32\u0e08\u0e40\u0e1b\u0e47\u0e19\u0e23\u0e2b\u0e31\u0e2a\u0e40\u0e02\u0e15\u0e2a\u0e38\u0e02\u0e20\u0e32\u0e1e\u0e17\u0e35\u0e48 6 \u0e41\u0e17\u0e19\u0e23\u0e2b\u0e31\u0e2a \u0e23\u0e1e.',
    ]
    for c in causes_2:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        run = p.add_run(c)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    add_styled_paragraph(doc, '\u0e27\u0e34\u0e18\u0e35\u0e41\u0e01\u0e49\u0e44\u0e02:',
                         font_size=14, bold=True, color=GREEN, space_before=8)
    fixes_2 = [
        '\u0e15\u0e23\u0e27\u0e08 Authen Code \u2014 \u0e15\u0e49\u0e2d\u0e07 valid \u0e13 \u0e27\u0e31\u0e19\u0e17\u0e35\u0e48 admit',
        '\u0e15\u0e23\u0e27\u0e08 IPD file: DATEADM, DATEDSC, DISCHT \u0e04\u0e23\u0e1a\u0e16\u0e39\u0e01\u0e15\u0e49\u0e2d\u0e07',
        '\u0e15\u0e23\u0e27\u0e08\u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23 Refer \u0e08\u0e32\u0e01 \u0e23\u0e1e.11855 \u2014 \u0e43\u0e1a\u0e2a\u0e48\u0e07\u0e15\u0e31\u0e27\u0e04\u0e23\u0e1a\u0e2b\u0e23\u0e37\u0e2d\u0e44\u0e21\u0e48',
        '\u0e40\u0e15\u0e23\u0e35\u0e22\u0e21 Cath report + clinical indication \u0e22\u0e37\u0e19\u0e22\u0e31\u0e19\u0e27\u0e48\u0e32\u0e15\u0e49\u0e2d\u0e07 admit \u0e08\u0e23\u0e34\u0e07',
    ]
    for f in fixes_2:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        run = p.add_run(f)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # --- Issue 3 ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('\u0e1b\u0e31\u0e0d\u0e2b\u0e32\u0e17\u0e35\u0e48 3: ')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = ORANGE
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    run2 = p.add_run('HC13 \u2014 \u0e22\u0e32 Clopidogrel (Deny Code: C)')
    run2.font.name = 'TH Sarabun New'
    run2.font.size = Pt(16)
    run2.bold = True
    run2.font.color.rgb = BLACK
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    issue3_table = doc.add_table(rows=3, cols=2)
    issue3_data = [
        ('\u0e01\u0e2d\u0e07\u0e17\u0e38\u0e19', 'HC13'),
        ('Sub-fund', 'CLOPIDOGREL_DRUG'),
        ('Deny Code', 'C'),
    ]
    for i, (label, value) in enumerate(issue3_data):
        cell_l = issue3_table.rows[i].cells[0]
        cell_v = issue3_table.rows[i].cells[1]
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        cell_l.width = Cm(4)
        p = cell_v.paragraphs[0]
        run = p.add_run(value)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        for cell in [cell_l, cell_v]:
            set_cell_border(cell,
                bottom={"val": "single", "sz": "2", "color": "E0E0E0"},
                top={"val": "nil", "sz": "0", "color": "FFFFFF"},
                left={"val": "nil", "sz": "0", "color": "FFFFFF"},
                right={"val": "nil", "sz": "0", "color": "FFFFFF"})

    add_styled_paragraph(doc, '\u0e2a\u0e32\u0e40\u0e2b\u0e15\u0e38\u0e17\u0e35\u0e48\u0e40\u0e1b\u0e47\u0e19\u0e44\u0e1b\u0e44\u0e14\u0e49:',
                         font_size=14, bold=True, color=DARK_RED, space_before=8)
    causes_3 = [
        'GPUID/TMT code \u0e44\u0e21\u0e48\u0e15\u0e23\u0e07 Drug Catalog \u0e02\u0e2d\u0e07 \u0e2a\u0e1b\u0e2a\u0e0a.',
        '\u0e08\u0e33\u0e19\u0e27\u0e19\u0e22\u0e32\u0e2b\u0e23\u0e37\u0e2d\u0e23\u0e32\u0e04\u0e32\u0e44\u0e21\u0e48\u0e15\u0e23\u0e07 Fee Schedule',
        '\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e43\u0e19 DRU file \u0e1c\u0e34\u0e14 format',
    ]
    for c in causes_3:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        run = p.add_run(c)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    add_styled_paragraph(doc, '\u0e27\u0e34\u0e18\u0e35\u0e41\u0e01\u0e49\u0e44\u0e02:',
                         font_size=14, bold=True, color=GREEN, space_before=8)
    fixes_3 = [
        '\u0e15\u0e23\u0e27\u0e08 DRU file: GPUID \u0e02\u0e2d\u0e07 Clopidogrel \u0e15\u0e23\u0e07 TMT catalog',
        '\u0e15\u0e23\u0e27\u0e08 Drug ID 24 \u0e2b\u0e25\u0e31\u0e01\u0e16\u0e39\u0e01\u0e15\u0e49\u0e2d\u0e07',
        '\u0e15\u0e23\u0e27\u0e08\u0e23\u0e32\u0e04\u0e32\u0e22\u0e32\u0e44\u0e21\u0e48\u0e40\u0e01\u0e34\u0e19 median price \u0e17\u0e35\u0e48\u0e01\u0e33\u0e2b\u0e19\u0e14',
        '\u0e41\u0e01\u0e49 code \u0e41\u0e25\u0e49\u0e27 resubmit',
    ]
    for f in fixes_3:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        run = p.add_run(f)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # ============================================================
    # SECTION 4: ACTION PLAN
    # ============================================================
    doc.add_page_break()
    add_section_header(doc, '4. \u0e41\u0e1c\u0e19\u0e01\u0e32\u0e23\u0e41\u0e01\u0e49\u0e44\u0e02 (Action Plan)')

    # Priority table
    action_table = doc.add_table(rows=4, cols=4)
    action_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['\u0e25\u0e33\u0e14\u0e31\u0e1a', '\u0e23\u0e32\u0e22\u0e01\u0e32\u0e23', '\u0e2a\u0e34\u0e48\u0e07\u0e17\u0e35\u0e48\u0e15\u0e49\u0e2d\u0e07\u0e17\u0e33', 'Impact']
    for j, h in enumerate(headers):
        cell = action_table.rows[0].cells[j]
        set_cell_shading(cell, '1A237E')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = WHITE
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    action_data = [
        ('[1]', '\u0e41\u0e01\u0e49 ADP file\n(INST \u2014 I04)',
         '\u0e15\u0e23\u0e27\u0e08 Device code, TYPE,\nSerial, \u0e23\u0e32\u0e04\u0e32', '\u0e2a\u0e39\u0e07\u0e2a\u0e38\u0e14'),
        ('[2]', '\u0e41\u0e01\u0e49 IP admission\n(IP01 \u2014 C)',
         '\u0e15\u0e23\u0e27\u0e08 Authen Code\n+ \u0e43\u0e1a Refer', '\u0e2a\u0e39\u0e07'),
        ('[3]', '\u0e41\u0e01\u0e49 Drug file\n(HC13 \u2014 Clopidogrel)',
         '\u0e15\u0e23\u0e27\u0e08 GPUID\n\u0e43\u0e19 DRU file', '\u0e1b\u0e32\u0e19\u0e01\u0e25\u0e32\u0e07'),
    ]

    bg_colors = ['FFEBEE', 'FFF3E0', 'FFF8E1']
    for i, (pri, item, action, impact) in enumerate(action_data):
        row = action_table.rows[i + 1]
        vals = [pri, item, action, impact]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            set_cell_shading(cell, bg_colors[i])
            p = cell.paragraphs[0]
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
            if j == 0:
                run.bold = True

    # Set column widths
    for row in action_table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(5.5)
        row.cells[3].width = Cm(3)

    # ============================================================
    # SECTION 5: RECOMMENDATION
    # ============================================================
    add_section_header(doc, '5. \u0e02\u0e49\u0e2d\u0e40\u0e2a\u0e19\u0e2d\u0e41\u0e19\u0e30 (Recommendation)')

    # Summary box
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="8" w:color="1A237E"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="8" w:color="1A237E"/>'
        f'<w:left w:val="single" w:sz="12" w:space="8" w:color="1A237E"/>'
        f'<w:right w:val="single" w:sz="12" w:space="8" w:color="1A237E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    # Add shading
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EAF6"/>')
    pPr.append(shd)

    run = p.add_run('\u0e21\u0e39\u0e25\u0e04\u0e48\u0e32\u0e40\u0e04\u0e2a: ~72,264 \u0e1a\u0e32\u0e17 \u2014 \u0e04\u0e38\u0e49\u0e21\u0e04\u0e48\u0e32\u0e21\u0e32\u0e01\u0e17\u0e35\u0e48\u0e08\u0e30\u0e41\u0e01\u0e49\u0e44\u0e02')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = NAVY
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    doc.add_paragraph()

    # Option A
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('\u0e17\u0e32\u0e07\u0e40\u0e25\u0e37\u0e2d\u0e01 A (\u0e41\u0e19\u0e30\u0e19\u0e33): ')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = GREEN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    run2 = p.add_run('\u0e41\u0e01\u0e49\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e43\u0e19 FDH \u0e41\u0e25\u0e49\u0e27 Resubmit')
    run2.font.name = 'TH Sarabun New'
    run2.font.size = Pt(14)
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    resubmit_steps = [
        '\u0e15\u0e23\u0e27\u0e08\u0e41\u0e25\u0e30\u0e41\u0e01\u0e49\u0e44\u0e02 ADP file (\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c) \u2014 \u0e43\u0e2a\u0e48 TYPE/CODE/Serial \u0e43\u0e2b\u0e49\u0e04\u0e23\u0e1a',
        '\u0e15\u0e23\u0e27\u0e08\u0e41\u0e25\u0e30\u0e41\u0e01\u0e49\u0e44\u0e02 IPD file + Authen Code + \u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23 Refer',
        '\u0e15\u0e23\u0e27\u0e08\u0e41\u0e25\u0e30\u0e41\u0e01\u0e49\u0e44\u0e02 DRU file (\u0e22\u0e32 Clopidogrel) \u2014 GPUID \u0e15\u0e23\u0e07 catalog',
        '\u0e2a\u0e48\u0e07\u0e40\u0e1a\u0e34\u0e01\u0e43\u0e2b\u0e21\u0e48\u0e1c\u0e48\u0e32\u0e19\u0e23\u0e30\u0e1a\u0e1a e-Claim',
    ]
    for s in resubmit_steps:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        run = p.add_run(s)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    doc.add_paragraph()

    # Option B
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('\u0e17\u0e32\u0e07\u0e40\u0e25\u0e37\u0e2d\u0e01 B: ')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = ORANGE
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    run2 = p.add_run('\u0e22\u0e37\u0e48\u0e19\u0e2d\u0e38\u0e17\u0e18\u0e23\u0e13\u0e4c\u0e1e\u0e23\u0e49\u0e2d\u0e21\u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23 clinical (\u0e16\u0e49\u0e32 resubmit \u0e44\u0e21\u0e48\u0e1c\u0e48\u0e32\u0e19)')
    run2.font.name = 'TH Sarabun New'
    run2.font.size = Pt(14)
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    appeal_steps = [
        '\u0e23\u0e48\u0e32\u0e07\u0e2b\u0e19\u0e31\u0e07\u0e2a\u0e37\u0e2d\u0e2d\u0e38\u0e17\u0e18\u0e23\u0e13\u0e4c\u0e16\u0e36\u0e07 \u0e2a\u0e1b\u0e2a\u0e0a.',
        '\u0e41\u0e19\u0e1a Cath report, EKG, Troponin results, \u0e43\u0e1a\u0e2a\u0e48\u0e07\u0e15\u0e31\u0e27',
        '\u0e2a\u0e48\u0e07\u0e20\u0e32\u0e22\u0e43\u0e19 30 \u0e27\u0e31\u0e19\u0e2b\u0e25\u0e31\u0e07\u0e44\u0e14\u0e49\u0e23\u0e31\u0e1a\u0e41\u0e08\u0e49\u0e07 deny',
    ]
    for s in appeal_steps:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.clear()
        run = p.add_run(s)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # ============================================================
    # SECTION 6: DRG DETAIL
    # ============================================================
    add_section_header(doc, '6. \u0e23\u0e32\u0e22\u0e25\u0e30\u0e40\u0e2d\u0e35\u0e22\u0e14 DRG')

    drg_info = [
        ('DRG Code', '05290'),
        ('MDC', '05 \u2014 Diseases & Disorders of the Circulatory System'),
        ('\u0e1b\u0e23\u0e30\u0e40\u0e20\u0e17', 'Surgical DRG (PCI/Stent)'),
        ('RW', '8.6544'),
        ('Base Rate (\u0e43\u0e19\u0e40\u0e02\u0e15)', '8,350 \u0e1a\u0e32\u0e17/AdjRW'),
        ('Base Rate (\u0e19\u0e2d\u0e01\u0e40\u0e02\u0e15)', '9,600 \u0e1a\u0e32\u0e17/AdjRW'),
        ('Expected Payment (\u0e43\u0e19\u0e40\u0e02\u0e15)', '72,264 \u0e1a\u0e32\u0e17'),
        ('Expected Payment (\u0e19\u0e2d\u0e01\u0e40\u0e02\u0e15)', '83,082 \u0e1a\u0e32\u0e17'),
        ('\u0e2b\u0e21\u0e32\u0e22\u0e40\u0e2b\u0e15\u0e38', 'RW \u0e2a\u0e39\u0e07 \u2192 \u0e21\u0e35 MCC \u0e23\u0e48\u0e27\u0e21\u0e14\u0e49\u0e27\u0e22 \u0e2b\u0e23\u0e37\u0e2d PCI \u0e1e\u0e23\u0e49\u0e2d\u0e21 DES stent'),
    ]
    create_info_table(doc, drg_info, col_widths=[Cm(5.5), Cm(10.5)])

    # ============================================================
    # SECTION 7: NOTE
    # ============================================================
    add_section_header(doc, '7. \u0e2b\u0e21\u0e32\u0e22\u0e40\u0e2b\u0e15\u0e38\u0e40\u0e1e\u0e34\u0e48\u0e21\u0e40\u0e15\u0e34\u0e21')

    notes = [
        '\u0e40\u0e04\u0e2a\u0e19\u0e35\u0e49\u0e40\u0e1b\u0e47\u0e19 Refer case (HMAIN \u2260 HCODE) \u0e15\u0e49\u0e2d\u0e07\u0e21\u0e35\u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23 Refer \u0e04\u0e23\u0e1a\u0e16\u0e49\u0e27\u0e19',
        'HMAIN2 = 06 \u0e1c\u0e34\u0e14\u0e1b\u0e01\u0e15\u0e34 \u2014 \u0e04\u0e27\u0e23\u0e15\u0e23\u0e27\u0e08\u0e01\u0e31\u0e1a \u0e2a\u0e1b\u0e2a\u0e0a. \u0e27\u0e48\u0e32\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e2a\u0e34\u0e17\u0e18\u0e34\u0e4c\u0e16\u0e39\u0e01\u0e15\u0e49\u0e2d\u0e07\u0e2b\u0e23\u0e37\u0e2d\u0e44\u0e21\u0e48',
        'DRG 05290 \u0e21\u0e35 RW \u0e2a\u0e39\u0e07 (8.6544) \u2014 \u0e04\u0e38\u0e49\u0e21\u0e04\u0e48\u0e32\u0e2d\u0e38\u0e17\u0e18\u0e23\u0e13\u0e4c/resubmit',
        '\u0e2b\u0e32\u0e01\u0e15\u0e49\u0e2d\u0e07\u0e2d\u0e38\u0e17\u0e18\u0e23\u0e13\u0e4c \u0e43\u0e2b\u0e49\u0e17\u0e33\u0e20\u0e32\u0e22\u0e43\u0e19 30 \u0e27\u0e31\u0e19\u0e2b\u0e25\u0e31\u0e07\u0e44\u0e14\u0e49\u0e23\u0e31\u0e1a\u0e41\u0e08\u0e49\u0e07 deny',
        '\u0e2b\u0e32\u0e01\u0e15\u0e49\u0e2d\u0e07\u0e01\u0e32\u0e23\u0e23\u0e48\u0e32\u0e07\u0e2b\u0e19\u0e31\u0e07\u0e2a\u0e37\u0e2d\u0e2d\u0e38\u0e17\u0e18\u0e23\u0e13\u0e4c \u0e2a\u0e32\u0e21\u0e32\u0e23\u0e16\u0e43\u0e0a\u0e49 Hospital Claim AI \u0e0a\u0e48\u0e27\u0e22\u0e23\u0e48\u0e32\u0e07\u0e44\u0e14\u0e49',
    ]
    for n in notes:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.clear()
        run = p.add_run(n)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(13)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # --- Final spacer + signature ---
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('\u0e23\u0e32\u0e22\u0e07\u0e32\u0e19\u0e2a\u0e23\u0e49\u0e32\u0e07\u0e42\u0e14\u0e22: Hospital Claim AI System')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = GRAY
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48: 31 \u0e21\u0e35\u0e19\u0e32\u0e04\u0e21 2569')
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = GRAY
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # === SAVE ===
    output_path = '/Users/ballbadboy/Desktop/projectX/Hospital claim AI/CathLab_Deny_Report_AN[REDACTED-AN].docx'
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    generate_report()
