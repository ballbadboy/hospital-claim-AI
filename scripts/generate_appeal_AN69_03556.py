#!/usr/bin/env python3
"""Generate formal appeal letter DOCX for AN 69-03556 Cath Lab Deny"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === Helpers ===
def set_thai_font(run, size=14, bold=False):
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

def add_thai_paragraph(doc, text, size=14, bold=False,
                       align=WD_ALIGN_PARAGRAPH.LEFT,
                       space_after=6, space_before=0,
                       first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_thai_font(run, size, bold)
    return p

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_cell_text(cell, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_thai_font(run, size, bold)

def generate_appeal():
    doc = Document()

    # --- Page Setup (A4, Thai government standard) ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # --- Default Style ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ==================================================================
    # HEADER: เลขที่หนังสือ + ชื่อหน่วยงาน
    # ==================================================================
    add_thai_paragraph(doc,
        '\u0e17\u0e35\u0e48 \u0e23\u0e1e. 10825/\u2026\u2026\u2026\u2026',
        size=16, bold=False, space_after=0)

    # ที่อยู่ (ชิดขวา)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        '\u0e42\u0e23\u0e07\u0e1e\u0e22\u0e32\u0e1a\u0e32\u0e25 (10825)\n'
        '\u0e08\u0e31\u0e07\u0e2b\u0e27\u0e31\u0e14\u0e19\u0e04\u0e23\u0e1b\u0e10\u0e21'
    )
    set_thai_font(run, 16)

    # วันที่ (ชิดขวา)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48 \u2026\u2026 \u0e40\u0e21\u0e29\u0e32\u0e22\u0e19 2569')
    set_thai_font(run, 16)

    # ==================================================================
    # เรื่อง
    # ==================================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('\u0e40\u0e23\u0e37\u0e48\u0e2d\u0e07  ')
    set_thai_font(run, 16, bold=True)
    run2 = p.add_run(
        '\u0e02\u0e2d\u0e17\u0e1a\u0e17\u0e27\u0e19\u0e1c\u0e25\u0e01\u0e32\u0e23\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a'
        '\u0e01\u0e32\u0e23\u0e02\u0e2d\u0e23\u0e31\u0e1a\u0e04\u0e48\u0e32\u0e43\u0e0a\u0e49\u0e08\u0e48\u0e32\u0e22'
        '\u0e40\u0e1e\u0e37\u0e48\u0e2d\u0e1a\u0e23\u0e34\u0e01\u0e32\u0e23\u0e2a\u0e32\u0e18\u0e32\u0e23\u0e13\u0e2a\u0e38\u0e02'
    )
    set_thai_font(run2, 16)

    # ==================================================================
    # เรียน
    # ==================================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('\u0e40\u0e23\u0e35\u0e22\u0e19  ')
    set_thai_font(run, 16, bold=True)
    run2 = p.add_run(
        '\u0e1c\u0e39\u0e49\u0e2d\u0e33\u0e19\u0e27\u0e22\u0e01\u0e32\u0e23'
        '\u0e2a\u0e33\u0e19\u0e31\u0e01\u0e07\u0e32\u0e19\u0e2b\u0e25\u0e31\u0e01\u0e1b\u0e23\u0e30\u0e01\u0e31\u0e19'
        '\u0e2a\u0e38\u0e02\u0e20\u0e32\u0e1e\u0e41\u0e2b\u0e48\u0e07\u0e0a\u0e32\u0e15\u0e34 '
        '\u0e40\u0e02\u0e15 6 \u0e23\u0e30\u0e22\u0e2d\u0e07'
    )
    set_thai_font(run2, 16)

    # ==================================================================
    # PARAGRAPH 1: อ้างอิง
    # ==================================================================
    add_thai_paragraph(doc,
        '\u0e15\u0e32\u0e21\u0e17\u0e35\u0e48\u0e42\u0e23\u0e07\u0e1e\u0e22\u0e32\u0e1a\u0e32\u0e25 (10825) '
        '\u0e23\u0e2b\u0e31\u0e2a\u0e2b\u0e19\u0e48\u0e27\u0e22\u0e1a\u0e23\u0e34\u0e01\u0e32\u0e23 10825 '
        '\u0e44\u0e14\u0e49\u0e2a\u0e48\u0e07\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e01\u0e32\u0e23\u0e02\u0e2d\u0e23\u0e31\u0e1a'
        '\u0e04\u0e48\u0e32\u0e43\u0e0a\u0e49\u0e08\u0e48\u0e32\u0e22\u0e40\u0e1e\u0e37\u0e48\u0e2d\u0e1a\u0e23\u0e34\u0e01\u0e32\u0e23'
        '\u0e2a\u0e32\u0e18\u0e32\u0e23\u0e13\u0e2a\u0e38\u0e02 '
        '\u0e2a\u0e33\u0e2b\u0e23\u0e31\u0e1a\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e23\u0e32\u0e22\u0e15\u0e48\u0e2d\u0e44\u0e1b\u0e19\u0e35\u0e49 '
        '\u0e1c\u0e48\u0e32\u0e19\u0e23\u0e30\u0e1a\u0e1a e-Claim \u0e41\u0e25\u0e49\u0e27 '
        '\u0e17\u0e32\u0e07\u0e42\u0e23\u0e07\u0e1e\u0e22\u0e32\u0e1a\u0e32\u0e25\u0e44\u0e14\u0e49\u0e23\u0e31\u0e1a\u0e41\u0e08\u0e49\u0e07\u0e27\u0e48\u0e32'
        '\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e14\u0e31\u0e07\u0e01\u0e25\u0e48\u0e32\u0e27'
        '\u0e44\u0e21\u0e48\u0e1c\u0e48\u0e32\u0e19\u0e01\u0e32\u0e23\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a '
        '\u0e17\u0e31\u0e49\u0e07\u0e19\u0e35\u0e49 '
        '\u0e02\u0e2d\u0e0a\u0e35\u0e49\u0e41\u0e08\u0e07\u0e14\u0e31\u0e07\u0e19\u0e35\u0e49',
        size=16, space_after=12, first_line_indent=2.5)

    # ==================================================================
    # PATIENT INFO TABLE
    # ==================================================================
    add_thai_paragraph(doc, '\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22',
                       size=16, bold=True, space_after=4, space_before=4)

    pt_table = doc.add_table(rows=12, cols=2)
    pt_table.style = 'Table Grid'
    pt_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    pt_data = [
        ('\u0e0a\u0e37\u0e48\u0e2d-\u0e2a\u0e01\u0e38\u0e25', '[\u0e0a\u0e37\u0e48\u0e2d\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22]'),
        ('\u0e40\u0e25\u0e02\u0e1b\u0e23\u0e30\u0e08\u0e33\u0e15\u0e31\u0e27\u0e1b\u0e23\u0e30\u0e0a\u0e32\u0e0a\u0e19 (CID)', '2411300021503'),
        ('HN', '69-09349'),
        ('AN', '69-03556'),
        ('\u0e2a\u0e34\u0e17\u0e18\u0e34\u0e4c\u0e01\u0e32\u0e23\u0e23\u0e31\u0e01\u0e29\u0e32', 'UC (\u0e2a\u0e34\u0e17\u0e18\u0e34\u0e4c\u0e2b\u0e25\u0e31\u0e01\u0e1b\u0e23\u0e30\u0e01\u0e31\u0e19\u0e2a\u0e38\u0e02\u0e20\u0e32\u0e1e\u0e16\u0e49\u0e27\u0e19\u0e2b\u0e19\u0e49\u0e32)'),
        ('\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48\u0e23\u0e31\u0e1a\u0e40\u0e02\u0e49\u0e32', '27 \u0e01\u0e38\u0e21\u0e20\u0e32\u0e1e\u0e31\u0e19\u0e18\u0e4c 2569 \u0e40\u0e27\u0e25\u0e32 11:44 \u0e19.'),
        ('\u0e27\u0e31\u0e19\u0e17\u0e35\u0e48\u0e08\u0e33\u0e2b\u0e19\u0e48\u0e32\u0e22', '1 \u0e21\u0e35\u0e19\u0e32\u0e04\u0e21 2569 \u0e40\u0e27\u0e25\u0e32 13:00 \u0e19.'),
        ('\u0e01\u0e32\u0e23\u0e27\u0e34\u0e19\u0e34\u0e08\u0e09\u0e31\u0e22\u0e2b\u0e25\u0e31\u0e01 (PDx)',
         '[\u0e23\u0e2b\u0e31\u0e2a ICD-10] [\u0e0a\u0e37\u0e48\u0e2d\u0e42\u0e23\u0e04] '
         '(\u0e40\u0e0a\u0e48\u0e19 I21.0 Acute transmural MI of anterior wall)'),
        ('\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23 (Procedures)',
         '[\u0e23\u0e2b\u0e31\u0e2a ICD-9-CM] [\u0e0a\u0e37\u0e48\u0e2d\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23] '
         '(\u0e40\u0e0a\u0e48\u0e19 36.07 PTCA with DES)'),
        ('DRG', '05290'),
        ('RW', '8.6544'),
        ('\u0e22\u0e2d\u0e14\u0e40\u0e1a\u0e34\u0e01 (Central Reimburse)', '71,322 \u0e1a\u0e32\u0e17'),
    ]
    for i, (label, value) in enumerate(pt_data):
        cell_l = pt_table.rows[i].cells[0]
        cell_v = pt_table.rows[i].cells[1]
        add_table_cell_text(cell_l, label, size=14, bold=True)
        add_table_cell_text(cell_v, value, size=14)
        cell_l.width = Cm(5)
        cell_v.width = Cm(10.5)
        set_cell_shading(cell_l, 'F5F5F5')

    # ==================================================================
    # DENY DETAILS TABLE
    # ==================================================================
    doc.add_paragraph()
    add_thai_paragraph(doc,
        '\u0e23\u0e32\u0e22\u0e25\u0e30\u0e40\u0e2d\u0e35\u0e22\u0e14\u0e01\u0e32\u0e23\u0e1b\u0e0f\u0e34\u0e40\u0e2a\u0e18'
        '\u0e01\u0e32\u0e23\u0e08\u0e48\u0e32\u0e22\u0e0a\u0e14\u0e40\u0e0a\u0e22',
        size=16, bold=True, space_after=4, space_before=8)

    deny_table = doc.add_table(rows=4, cols=4)
    deny_table.style = 'Table Grid'
    deny_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    deny_headers = [
        '\u0e23\u0e2b\u0e31\u0e2a Deny',
        '\u0e04\u0e27\u0e32\u0e21\u0e2b\u0e21\u0e32\u0e22',
        '\u0e23\u0e32\u0e22\u0e01\u0e32\u0e23\u0e17\u0e35\u0e48\u0e16\u0e39\u0e01\u0e1b\u0e0f\u0e34\u0e40\u0e2a\u0e18',
        '\u0e21\u0e39\u0e25\u0e04\u0e48\u0e32\u0e42\u0e14\u0e22\u0e1b\u0e23\u0e30\u0e21\u0e32\u0e13',
    ]
    for j, h in enumerate(deny_headers):
        cell = deny_table.rows[0].cells[j]
        set_cell_shading(cell, '1A237E')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_thai_font(run, 13, bold=True)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    deny_data = [
        ('HC09 / I04',
         '\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e41\u0e25\u0e30\u0e2d\u0e27\u0e31\u0e22\u0e27\u0e30\u0e40\u0e17\u0e35\u0e22\u0e21\n\u0e43\u0e19\u0e01\u0e32\u0e23\u0e17\u0e33\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23',
         'INST (Stent/\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e2a\u0e27\u0e19\u0e2b\u0e31\u0e27\u0e43\u0e08)',
         '\u0e1b\u0e23\u0e30\u0e21\u0e32\u0e13 31,490 \u0e1a\u0e32\u0e17'),
        ('IP01 / C',
         '\u0e40\u0e01\u0e13\u0e11\u0e4c\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e43\u0e19',
         'IPINRGR (IP \u0e43\u0e19\u0e40\u0e02\u0e15)',
         '\u0e1b\u0e23\u0e30\u0e21\u0e32\u0e13 40,000 \u0e1a\u0e32\u0e17'),
        ('HC13 / C',
         '\u0e22\u0e32 Clopidogrel',
         'CLOPIDOGREL_DRUG',
         '\u0e15\u0e32\u0e21 Fee Schedule'),
    ]
    for i, (code, meaning, item, value) in enumerate(deny_data):
        row = deny_table.rows[i + 1]
        vals = [code, meaning, item, value]
        for j, val in enumerate(vals):
            add_table_cell_text(row.cells[j], val, size=13)
        set_cell_shading(row.cells[0], 'FFEBEE')

    # Set column widths
    for row in deny_table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(5)
        row.cells[3].width = Cm(4)

    # ==================================================================
    # ข้อชี้แจง (CLINICAL JUSTIFICATION)
    # ==================================================================
    doc.add_paragraph()
    add_thai_paragraph(doc,
        '\u0e17\u0e31\u0e49\u0e07\u0e19\u0e35\u0e49 \u0e02\u0e2d\u0e0a\u0e35\u0e49\u0e41\u0e08\u0e07\u0e14\u0e31\u0e07\u0e19\u0e35\u0e49',
        size=16, bold=True, space_after=8, space_before=12)

    # --- ข้อ 1: ความจำเป็นทาง clinical ---
    add_thai_paragraph(doc,
        '1. \u0e04\u0e27\u0e32\u0e21\u0e08\u0e33\u0e40\u0e1b\u0e47\u0e19\u0e17\u0e32\u0e07\u0e04\u0e25\u0e34\u0e19\u0e34\u0e01'
        '\u0e43\u0e19\u0e01\u0e32\u0e23\u0e23\u0e31\u0e1a\u0e40\u0e02\u0e49\u0e32\u0e23\u0e31\u0e01\u0e29\u0e32\u0e40\u0e1b\u0e47\u0e19'
        '\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e43\u0e19 (IP01)',
        size=16, bold=True, space_after=4, space_before=8)

    add_thai_paragraph(doc,
        '\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e23\u0e32\u0e22\u0e19\u0e35\u0e49'
        '\u0e40\u0e02\u0e49\u0e32\u0e23\u0e31\u0e1a\u0e01\u0e32\u0e23\u0e23\u0e31\u0e01\u0e29\u0e32\u0e14\u0e49\u0e27\u0e22'
        '\u0e2d\u0e32\u0e01\u0e32\u0e23 [\u0e23\u0e30\u0e1a\u0e38\u0e2d\u0e32\u0e01\u0e32\u0e23: \u0e40\u0e0a\u0e48\u0e19 '
        '\u0e40\u0e08\u0e47\u0e1a\u0e41\u0e19\u0e48\u0e19\u0e2b\u0e19\u0e49\u0e32\u0e2d\u0e01, '
        '\u0e2b\u0e32\u0e22\u0e43\u0e08\u0e02\u0e13\u0e30\u0e1e\u0e31\u0e01] '
        '\u0e15\u0e23\u0e27\u0e08\u0e1e\u0e1a EKG \u0e41\u0e2a\u0e14\u0e07 '
        '[\u0e23\u0e30\u0e1a\u0e38\u0e1c\u0e25 EKG: \u0e40\u0e0a\u0e48\u0e19 '
        'ST elevation \u0e43\u0e19 lead V1-V4] '
        '\u0e1c\u0e25 Troponin [\u0e23\u0e30\u0e1a\u0e38\u0e04\u0e48\u0e32] '
        '\u0e0b\u0e36\u0e48\u0e07\u0e40\u0e02\u0e49\u0e32\u0e01\u0e31\u0e1a\u0e01\u0e32\u0e23\u0e27\u0e34\u0e19\u0e34\u0e08\u0e09\u0e31\u0e22 '
        '[\u0e23\u0e30\u0e1a\u0e38 diagnosis \u0e20\u0e32\u0e29\u0e32\u0e2d\u0e31\u0e07\u0e01\u0e24\u0e29] '
        '\u0e08\u0e36\u0e07\u0e21\u0e35\u0e04\u0e27\u0e32\u0e21\u0e08\u0e33\u0e40\u0e1b\u0e47\u0e19\u0e15\u0e49\u0e2d\u0e07'
        '\u0e23\u0e31\u0e1a\u0e44\u0e27\u0e49\u0e40\u0e1b\u0e47\u0e19\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e43\u0e19'
        '\u0e40\u0e1e\u0e37\u0e48\u0e2d\u0e40\u0e15\u0e23\u0e35\u0e22\u0e21\u0e17\u0e33\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23'
        '\u0e2a\u0e27\u0e19\u0e2b\u0e31\u0e27\u0e43\u0e08\u0e41\u0e25\u0e30\u0e14\u0e39\u0e41\u0e25\u0e2b\u0e25\u0e31\u0e07\u0e17\u0e33'
        '\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23 '
        '\u0e14\u0e31\u0e07\u0e1b\u0e23\u0e32\u0e01\u0e0f\u0e15\u0e32\u0e21\u0e40\u0e27\u0e0a\u0e23\u0e30\u0e40\u0e1a\u0e35\u0e22\u0e19\u0e41\u0e19\u0e1a',
        size=16, space_after=8, first_line_indent=2.5)

    add_thai_paragraph(doc,
        '\u0e01\u0e32\u0e23\u0e23\u0e31\u0e1a\u0e44\u0e27\u0e49\u0e40\u0e1b\u0e47\u0e19\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e43\u0e19'
        '\u0e40\u0e1b\u0e47\u0e19\u0e44\u0e1b\u0e15\u0e32\u0e21\u0e41\u0e19\u0e27\u0e17\u0e32\u0e07\u0e40\u0e27\u0e0a\u0e1b\u0e0f\u0e34\u0e1a\u0e31\u0e15\u0e34 '
        '(Clinical Practice Guideline) '
        '\u0e02\u0e2d\u0e07\u0e2a\u0e21\u0e32\u0e04\u0e21\u0e42\u0e23\u0e04\u0e2b\u0e31\u0e27\u0e43\u0e08\u0e41\u0e2b\u0e48\u0e07\u0e1b\u0e23\u0e30\u0e40\u0e17\u0e28\u0e44\u0e17\u0e22 '
        '\u0e41\u0e25\u0e30 ACC/AHA Guidelines '
        '\u0e0b\u0e36\u0e48\u0e07\u0e23\u0e30\u0e1a\u0e38\u0e27\u0e48\u0e32\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22'
        '\u0e17\u0e35\u0e48\u0e21\u0e35\u0e2d\u0e32\u0e01\u0e32\u0e23\u0e40\u0e09\u0e35\u0e22\u0e1a\u0e1e\u0e25\u0e31\u0e19'
        '\u0e41\u0e25\u0e30\u0e1c\u0e25\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e19\u0e31\u0e1a\u0e2a\u0e19\u0e38\u0e19 '
        '\u0e21\u0e35\u0e02\u0e49\u0e2d\u0e1a\u0e48\u0e07\u0e0a\u0e35\u0e49\u0e43\u0e19\u0e01\u0e32\u0e23\u0e17\u0e33 '
        'Cardiac Catheterization \u0e41\u0e25\u0e30/\u0e2b\u0e23\u0e37\u0e2d '
        'Percutaneous Coronary Intervention (PCI) '
        '\u0e42\u0e14\u0e22\u0e15\u0e49\u0e2d\u0e07\u0e23\u0e31\u0e1a\u0e44\u0e27\u0e49\u0e40\u0e1b\u0e47\u0e19\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e43\u0e19'
        '\u0e40\u0e1e\u0e37\u0e48\u0e2d\u0e40\u0e1d\u0e49\u0e32\u0e23\u0e30\u0e27\u0e31\u0e07\u0e20\u0e32\u0e27\u0e30\u0e41\u0e17\u0e23\u0e01\u0e0b\u0e49\u0e2d\u0e19',
        size=16, space_after=8, first_line_indent=2.5)

    # --- ข้อ 2: อุปกรณ์ HC09 ---
    add_thai_paragraph(doc,
        '2. \u0e01\u0e23\u0e13\u0e35\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e41\u0e25\u0e30\u0e2d\u0e27\u0e31\u0e22\u0e27\u0e30\u0e40\u0e17\u0e35\u0e22\u0e21 (HC09/I04)',
        size=16, bold=True, space_after=4, space_before=8)

    add_thai_paragraph(doc,
        '\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e17\u0e35\u0e48\u0e43\u0e0a\u0e49\u0e43\u0e19\u0e01\u0e32\u0e23\u0e17\u0e33'
        '\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23\u0e04\u0e23\u0e31\u0e49\u0e07\u0e19\u0e35\u0e49\u0e40\u0e1b\u0e47\u0e19 '
        '[\u0e23\u0e30\u0e1a\u0e38\u0e1b\u0e23\u0e30\u0e40\u0e20\u0e17 stent: BMS/DES] '
        '\u0e22\u0e35\u0e48\u0e2b\u0e49\u0e2d [\u0e23\u0e30\u0e1a\u0e38\u0e22\u0e35\u0e48\u0e2b\u0e49\u0e2d] '
        '\u0e0b\u0e36\u0e48\u0e07\u0e40\u0e1b\u0e47\u0e19\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e17\u0e35\u0e48'
        '\u0e2d\u0e22\u0e39\u0e48\u0e43\u0e19\u0e1a\u0e31\u0e0d\u0e0a\u0e35'
        '\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e17\u0e35\u0e48 \u0e2a\u0e1b\u0e2a\u0e0a. \u0e2d\u0e19\u0e38\u0e21\u0e31\u0e15\u0e34 '
        '\u0e42\u0e14\u0e22\u0e21\u0e35 Lot Number [\u0e23\u0e30\u0e1a\u0e38] '
        '\u0e41\u0e25\u0e30 Serial Number [\u0e23\u0e30\u0e1a\u0e38] '
        '\u0e17\u0e31\u0e49\u0e07\u0e19\u0e35\u0e49 \u0e01\u0e32\u0e23\u0e17\u0e35\u0e48\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e44\u0e21\u0e48\u0e1c\u0e48\u0e32\u0e19'
        '\u0e2d\u0e32\u0e08\u0e40\u0e01\u0e34\u0e14\u0e08\u0e32\u0e01\u0e01\u0e32\u0e23\u0e1a\u0e31\u0e19\u0e17\u0e36\u0e01\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25'
        '\u0e43\u0e19 ADP file \u0e44\u0e21\u0e48\u0e04\u0e23\u0e1a\u0e16\u0e49\u0e27\u0e19 '
        '\u0e0b\u0e36\u0e48\u0e07\u0e44\u0e14\u0e49\u0e14\u0e33\u0e40\u0e19\u0e34\u0e19\u0e01\u0e32\u0e23\u0e41\u0e01\u0e49\u0e44\u0e02\u0e41\u0e25\u0e49\u0e27 '
        '\u0e14\u0e31\u0e07\u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23\u0e41\u0e19\u0e1a',
        size=16, space_after=8, first_line_indent=2.5)

    # --- ข้อ 3: ยา Clopidogrel HC13 ---
    add_thai_paragraph(doc,
        '3. \u0e01\u0e23\u0e13\u0e35\u0e22\u0e32 Clopidogrel (HC13)',
        size=16, bold=True, space_after=4, space_before=8)

    add_thai_paragraph(doc,
        '\u0e22\u0e32 Clopidogrel \u0e40\u0e1b\u0e47\u0e19\u0e22\u0e32\u0e15\u0e49\u0e32\u0e19\u0e40\u0e01\u0e25\u0e47\u0e14\u0e40\u0e25\u0e37\u0e2d\u0e14 '
        '(Antiplatelet) \u0e17\u0e35\u0e48\u0e08\u0e33\u0e40\u0e1b\u0e47\u0e19\u0e2b\u0e25\u0e31\u0e07\u0e17\u0e33 PCI '
        '\u0e15\u0e32\u0e21\u0e41\u0e19\u0e27\u0e17\u0e32\u0e07\u0e40\u0e27\u0e0a\u0e1b\u0e0f\u0e34\u0e1a\u0e31\u0e15\u0e34 '
        'ACC/AHA \u0e41\u0e25\u0e30\u0e2a\u0e21\u0e32\u0e04\u0e21\u0e42\u0e23\u0e04\u0e2b\u0e31\u0e27\u0e43\u0e08\u0e41\u0e2b\u0e48\u0e07\u0e1b\u0e23\u0e30\u0e40\u0e17\u0e28\u0e44\u0e17\u0e22 '
        '\u0e42\u0e14\u0e22\u0e15\u0e49\u0e2d\u0e07\u0e43\u0e2b\u0e49 Dual Antiplatelet Therapy (DAPT) '
        '\u0e2b\u0e25\u0e31\u0e07\u0e43\u0e2a\u0e48 stent '
        '\u0e17\u0e31\u0e49\u0e07\u0e19\u0e35\u0e49 \u0e01\u0e32\u0e23\u0e17\u0e35\u0e48\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e22\u0e32\u0e44\u0e21\u0e48\u0e1c\u0e48\u0e32\u0e19'
        '\u0e2d\u0e32\u0e08\u0e40\u0e01\u0e34\u0e14\u0e08\u0e32\u0e01 GPUID/TMT code \u0e44\u0e21\u0e48\u0e15\u0e23\u0e07\u0e01\u0e31\u0e1a Drug Catalog '
        '\u0e0b\u0e36\u0e48\u0e07\u0e44\u0e14\u0e49\u0e14\u0e33\u0e40\u0e19\u0e34\u0e19\u0e01\u0e32\u0e23\u0e41\u0e01\u0e49\u0e44\u0e02\u0e43\u0e19 DRU file \u0e41\u0e25\u0e49\u0e27',
        size=16, space_after=8, first_line_indent=2.5)

    # --- ข้อ 4: ข้อมูลที่แก้ไข ---
    add_thai_paragraph(doc,
        '4. \u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e17\u0e35\u0e48\u0e44\u0e14\u0e49\u0e14\u0e33\u0e40\u0e19\u0e34\u0e19\u0e01\u0e32\u0e23\u0e41\u0e01\u0e49\u0e44\u0e02',
        size=16, bold=True, space_after=4, space_before=8)

    add_thai_paragraph(doc,
        '\u0e17\u0e32\u0e07\u0e42\u0e23\u0e07\u0e1e\u0e22\u0e32\u0e1a\u0e32\u0e25\u0e44\u0e14\u0e49\u0e14\u0e33\u0e40\u0e19\u0e34\u0e19\u0e01\u0e32\u0e23'
        '\u0e41\u0e01\u0e49\u0e44\u0e02\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e14\u0e31\u0e07\u0e15\u0e48\u0e2d\u0e44\u0e1b\u0e19\u0e35\u0e49:',
        size=16, space_after=4, first_line_indent=2.5)

    corrections = [
        '\u0e41\u0e01\u0e49\u0e44\u0e02 ADP file: \u0e40\u0e1e\u0e34\u0e48\u0e21 TYPE, CODE, SERIALNO \u0e02\u0e2d\u0e07\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e43\u0e2b\u0e49\u0e04\u0e23\u0e1a\u0e16\u0e49\u0e27\u0e19\u0e15\u0e32\u0e21 NHSO Device Catalog',
        '\u0e41\u0e01\u0e49\u0e44\u0e02 DRU file: \u0e1b\u0e23\u0e31\u0e1a GPUID \u0e02\u0e2d\u0e07 Clopidogrel \u0e43\u0e2b\u0e49\u0e15\u0e23\u0e07\u0e01\u0e31\u0e1a TMT Drug Catalog',
        '\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a IPD file: DATEADM, DATEDSC, DISCHT \u0e04\u0e23\u0e1a\u0e16\u0e49\u0e27\u0e19',
        '\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a Authen Code \u0e41\u0e25\u0e30\u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23\u0e01\u0e32\u0e23\u0e2a\u0e48\u0e07\u0e15\u0e31\u0e27 (Refer) \u0e04\u0e23\u0e1a\u0e16\u0e49\u0e27\u0e19',
    ]
    for c in corrections:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(3)
        p.clear()
        run = p.add_run(c)
        set_thai_font(run, 16)

    # ==================================================================
    # อ้างอิงหลักเกณฑ์
    # ==================================================================
    doc.add_paragraph()
    add_thai_paragraph(doc,
        '\u0e15\u0e32\u0e21\u0e1b\u0e23\u0e30\u0e01\u0e32\u0e28\u0e2a\u0e33\u0e19\u0e31\u0e01\u0e07\u0e32\u0e19'
        '\u0e2b\u0e25\u0e31\u0e01\u0e1b\u0e23\u0e30\u0e01\u0e31\u0e19\u0e2a\u0e38\u0e02\u0e20\u0e32\u0e1e\u0e41\u0e2b\u0e48\u0e07\u0e0a\u0e32\u0e15\u0e34 '
        '\u0e40\u0e23\u0e37\u0e48\u0e2d\u0e07 '
        '\u0e2b\u0e25\u0e31\u0e01\u0e40\u0e01\u0e13\u0e11\u0e4c\u0e01\u0e32\u0e23\u0e14\u0e33\u0e40\u0e19\u0e34\u0e19\u0e07\u0e32\u0e19'
        '\u0e41\u0e25\u0e30\u0e01\u0e32\u0e23\u0e1a\u0e23\u0e34\u0e2b\u0e32\u0e23\u0e01\u0e2d\u0e07\u0e17\u0e38\u0e19'
        '\u0e2b\u0e25\u0e31\u0e01\u0e1b\u0e23\u0e30\u0e01\u0e31\u0e19\u0e2a\u0e38\u0e02\u0e20\u0e32\u0e1e\u0e41\u0e2b\u0e48\u0e07\u0e0a\u0e32\u0e15\u0e34 '
        '\u0e1b\u0e35\u0e07\u0e1a\u0e1b\u0e23\u0e30\u0e21\u0e32\u0e13 2569 '
        '\u0e02\u0e49\u0e2d 8 \u0e01\u0e33\u0e2b\u0e19\u0e14\u0e27\u0e48\u0e32 '
        '\u0e2b\u0e19\u0e48\u0e27\u0e22\u0e1a\u0e23\u0e34\u0e01\u0e32\u0e23\u0e21\u0e35\u0e2a\u0e34\u0e17\u0e18\u0e34\u0e4c'
        '\u0e02\u0e2d\u0e23\u0e31\u0e1a\u0e04\u0e48\u0e32\u0e43\u0e0a\u0e49\u0e08\u0e48\u0e32\u0e22\u0e2a\u0e33\u0e2b\u0e23\u0e31\u0e1a'
        '\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c\u0e41\u0e25\u0e30\u0e2d\u0e27\u0e31\u0e22\u0e27\u0e30\u0e40\u0e17\u0e35\u0e22\u0e21'
        '\u0e43\u0e19\u0e01\u0e32\u0e23\u0e17\u0e33\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23'
        '\u0e17\u0e35\u0e48\u0e21\u0e35\u0e04\u0e27\u0e32\u0e21\u0e08\u0e33\u0e40\u0e1b\u0e47\u0e19\u0e17\u0e32\u0e07\u0e01\u0e32\u0e23\u0e41\u0e1e\u0e17\u0e22\u0e4c '
        '\u0e41\u0e25\u0e30\u0e04\u0e48\u0e32\u0e22\u0e32\u0e17\u0e35\u0e48\u0e08\u0e33\u0e40\u0e1b\u0e47\u0e19'
        '\u0e2b\u0e25\u0e31\u0e07\u0e17\u0e33\u0e2b\u0e31\u0e15\u0e16\u0e01\u0e32\u0e23\u0e2a\u0e27\u0e19\u0e2b\u0e31\u0e27\u0e43\u0e08',
        size=16, space_after=12, first_line_indent=2.5)

    # ==================================================================
    # CLOSING
    # ==================================================================
    add_thai_paragraph(doc,
        '\u0e08\u0e36\u0e07\u0e40\u0e23\u0e35\u0e22\u0e19\u0e21\u0e32\u0e40\u0e1e\u0e37\u0e48\u0e2d'
        '\u0e42\u0e1b\u0e23\u0e14\u0e1e\u0e34\u0e08\u0e32\u0e23\u0e13\u0e32\u0e17\u0e1a\u0e17\u0e27\u0e19'
        '\u0e1c\u0e25\u0e01\u0e32\u0e23\u0e15\u0e23\u0e27\u0e08\u0e2a\u0e2d\u0e1a '
        '\u0e41\u0e25\u0e30\u0e02\u0e2d\u0e43\u0e2b\u0e49\u0e14\u0e33\u0e40\u0e19\u0e34\u0e19\u0e01\u0e32\u0e23'
        '\u0e08\u0e48\u0e32\u0e22\u0e0a\u0e14\u0e40\u0e0a\u0e22\u0e04\u0e48\u0e32\u0e1a\u0e23\u0e34\u0e01\u0e32\u0e23'
        '\u0e15\u0e32\u0e21\u0e2b\u0e25\u0e31\u0e01\u0e40\u0e01\u0e13\u0e11\u0e4c\u0e17\u0e35\u0e48\u0e01\u0e33\u0e2b\u0e19\u0e14 '
        '\u0e08\u0e30\u0e40\u0e1b\u0e47\u0e19\u0e1e\u0e23\u0e30\u0e04\u0e38\u0e13\u0e22\u0e34\u0e48\u0e07',
        size=16, space_after=24, space_before=12, first_line_indent=2.5)

    # --- ลงนาม ---
    add_thai_paragraph(doc,
        '\u0e02\u0e2d\u0e41\u0e2a\u0e14\u0e07\u0e04\u0e27\u0e32\u0e21\u0e19\u0e31\u0e1a\u0e16\u0e37\u0e2d',
        size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36, space_before=12)

    add_thai_paragraph(doc,
        '(\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026)',
        size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_thai_paragraph(doc,
        '\u0e15\u0e33\u0e41\u0e2b\u0e19\u0e48\u0e07 \u0e1c\u0e39\u0e49\u0e2d\u0e33\u0e19\u0e27\u0e22\u0e01\u0e32\u0e23\u0e42\u0e23\u0e07\u0e1e\u0e22\u0e32\u0e1a\u0e32\u0e25',
        size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    # ==================================================================
    # เอกสารแนบ
    # ==================================================================
    doc.add_paragraph()
    add_thai_paragraph(doc,
        '\u0e40\u0e2d\u0e01\u0e2a\u0e32\u0e23\u0e41\u0e19\u0e1a:',
        size=16, bold=True, space_after=4, space_before=12)

    attachments = [
        '\u0e2a\u0e33\u0e40\u0e19\u0e32\u0e40\u0e27\u0e0a\u0e23\u0e30\u0e40\u0e1a\u0e35\u0e22\u0e19\u0e1c\u0e39\u0e49\u0e1b\u0e48\u0e27\u0e22\u0e43\u0e19 (IPD)',
        '\u0e23\u0e32\u0e22\u0e07\u0e32\u0e19\u0e1c\u0e25\u0e01\u0e32\u0e23\u0e2a\u0e27\u0e19\u0e2b\u0e31\u0e27\u0e43\u0e08 (Cardiac Catheterization Report)',
        '\u0e1c\u0e25 EKG',
        '\u0e1c\u0e25\u0e15\u0e23\u0e27\u0e08 Troponin / Cardiac Enzymes',
        '\u0e1c\u0e25 Echocardiography (\u0e16\u0e49\u0e32\u0e21\u0e35)',
        '\u0e43\u0e1a\u0e2a\u0e48\u0e07\u0e15\u0e31\u0e27\u0e08\u0e32\u0e01 \u0e23\u0e1e.11855 (Refer document)',
        'Authen Code / \u0e2b\u0e25\u0e31\u0e01\u0e10\u0e32\u0e19\u0e2a\u0e34\u0e17\u0e18\u0e34\u0e4c\u0e01\u0e32\u0e23\u0e23\u0e31\u0e01\u0e29\u0e32',
        '\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e2d\u0e38\u0e1b\u0e01\u0e23\u0e13\u0e4c (Stent: Lot/Serial, Invoice)',
        '\u0e02\u0e49\u0e2d\u0e21\u0e39\u0e25\u0e22\u0e32 Clopidogrel (GPUID, Drug Invoice)',
    ]
    for i, att in enumerate(attachments, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(2)
        run = p.add_run(f'{i}. {att}')
        set_thai_font(run, 16)

    # ==================================================================
    # SAVE
    # ==================================================================
    output_path = '/Users/ballbadboy/Desktop/projectX/Hospital claim AI/Appeal_AN69-03556.docx'
    doc.save(output_path)
    print(f'Appeal letter saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    generate_appeal()
