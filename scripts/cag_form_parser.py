#!/usr/bin/env python3
"""
CAG Data Form Parser — อ่าน CAG form (DOCX) ที่กรอกแล้ว ดึงข้อมูลอัตโนมัติ

Usage:
    from scripts.cag_form_parser import parse_cag_form
    data = parse_cag_form("path/to/filled_cag_form.docx")

Output: dict with structured fields for use in V993 appeal letters
"""

import re
from pathlib import Path
from docx import Document


# ─── Checkbox Detection ─────────────────────────────────

# Unicode checkboxes used in Thai hospital forms
CHECKED_CHARS = {"☑", "☒", "■", "◼", "✓", "✔", "⬛"}
UNCHECKED_CHARS = {"☐", "□", "◻", "⬜"}


def _is_checked(text):
    """Check if a checkbox text indicates 'checked'."""
    for ch in CHECKED_CHARS:
        if ch in text:
            return True
    return False


def _extract_after(text, keyword):
    """Extract text after a keyword, stripping dots and whitespace."""
    idx = text.find(keyword)
    if idx < 0:
        return ""
    val = text[idx + len(keyword):]
    val = re.sub(r'[…\.]+', '', val).strip()
    return val


def _extract_number(text, keyword):
    """Extract a number after a keyword."""
    val = _extract_after(text, keyword)
    m = re.search(r'[\d.]+', val)
    return float(m.group()) if m else None


def _find_checked_option(text, options):
    """Find which option is checked in a line with multiple checkboxes."""
    for opt in options:
        # Look for checked char near the option text
        pattern = rf'[{"".join(re.escape(c) for c in CHECKED_CHARS)}]\s*{re.escape(opt)}'
        if re.search(pattern, text):
            return opt
    return ""


# ─── Main Parser ────────────────────────────────────────

def parse_cag_form(docx_path):
    """Parse filled CAG Data Form DOCX → structured dict."""
    doc = Document(docx_path)

    # Collect all text (paragraphs + tables)
    all_lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            all_lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    all_lines.append(t)

    full_text = "\n".join(all_lines)

    data = {
        "general": _parse_general(all_lines, full_text),
        "risk_factors": _parse_risk_factors(all_lines, full_text),
        "indication": _parse_indication(all_lines, full_text),
        "cag_procedure": _parse_cag_procedure(all_lines, full_text),
        "pci_procedure": _parse_pci_procedure(all_lines, full_text),
    }

    return data


def _parse_general(lines, full_text):
    """Section A: General Information."""
    result = {
        "hcode": "",
        "pid": "",
        "sex": "",
        "admission_date": "",
    }

    for line in lines:
        if "HCODE" in line:
            m = re.search(r'HCODE\s*(\d{5})', line)
            if m:
                result["hcode"] = m.group(1)

        if "PID" in line:
            m = re.search(r'PID\s*(\d{13})', line)
            if m:
                result["pid"] = m.group(1)

        if "Admission Date" in line or "Admission date" in line:
            m = re.search(r'(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})', line)
            if m:
                result["admission_date"] = m.group(1)

        # Sex
        if "ชาย" in line or "หญิง" in line:
            if _find_checked_option(line, ["ชาย"]):
                result["sex"] = "ชาย"
            elif _find_checked_option(line, ["หญิง"]):
                result["sex"] = "หญิง"

    return result


def _parse_risk_factors(lines, full_text):
    """Section B: History and Risk Factors."""
    result = {
        "height_cm": None,
        "weight_kg": None,
        "bmi": None,
        "hypertension": None,
        "dyslipidemia": None,
        "ldl": None,
        "hdl": None,
        "dm": None,
        "smoker": None,
        "prior_mi": None,
        "cerebrovascular_disease": None,
        "peripheral_arterial_disease": None,
    }

    yes_no_fields = {
        "Hypertension": "hypertension",
        "Dyslipidemia": "dyslipidemia",
        "DM": "dm",
        "Prior MI": "prior_mi",
        "smoker": "smoker",
        "Cerebrovascular": "cerebrovascular_disease",
        "Peripheral arterial": "peripheral_arterial_disease",
    }

    for line in lines:
        # Height/Weight/BMI
        if "Height" in line:
            result["height_cm"] = _extract_number(line, "Height")
        if "Weight" in line:
            result["weight_kg"] = _extract_number(line, "Weight")
        if "BMI" in line:
            result["bmi"] = _extract_number(line, "BMI")
        if "LDL" in line:
            result["ldl"] = _extract_number(line, "LDL")
        if "HDL" in line:
            result["hdl"] = _extract_number(line, "HDL")

        # Yes/No fields
        for keyword, field in yes_no_fields.items():
            if keyword.lower() in line.lower():
                if _find_checked_option(line, ["Yes"]):
                    result[field] = True
                elif _find_checked_option(line, ["No"]):
                    result[field] = False

    return result


def _parse_indication(lines, full_text):
    """Section C: Clinical Indication."""
    result = {
        "type": "",  # C1/C2/C3/C4/C5/C6
        "lvef": "",
        "lvef_value": None,
        "grace_score": None,
        "risk_level": "",  # very_high/high/intermediate/low
        "shock": False,
        "vt_vf": False,
    }

    # Detect case type from section headers
    for line in lines:
        line_lower = line.lower()

        # Case type detection
        if "c1" in line_lower and ("ccs" in line_lower or "chronic" in line_lower):
            result["type"] = "C1"
        elif "c2" in line_lower and ("nste" in line_lower or "non-st" in line_lower):
            result["type"] = "C2"
        elif "c3" in line_lower and ("ste-acs" in line_lower or "stemi" in line_lower or "st elevation" in line_lower):
            result["type"] = "C3"
        elif "c4" in line_lower and ("cmp" in line_lower or "cardiomyopathy" in line_lower):
            result["type"] = "C4"
        elif "c5" in line_lower and "pre-op" in line_lower:
            result["type"] = "C5"
        elif "c6" in line_lower and "cardiac arrest" in line_lower:
            result["type"] = "C6"

        # LVEF
        if "LVEF" in line or "EF" in line:
            if _find_checked_option(line, ["<35%"]):
                result["lvef"] = "<35%"
                result["lvef_value"] = 30  # estimate
            elif _find_checked_option(line, ["35-50%", "36-50%"]):
                result["lvef"] = "35-50%"
                result["lvef_value"] = 42
            elif _find_checked_option(line, [">50%"]):
                result["lvef"] = ">50%"
                result["lvef_value"] = 55

        # Risk level
        if "very high" in line_lower:
            if _is_checked(line):
                result["risk_level"] = "very_high"
        elif "high" in line_lower and "grace" in line_lower and "> 140" in line:
            if _is_checked(line):
                result["risk_level"] = "high"
                result["grace_score"] = 141
        elif "intermediate" in line_lower or "low" in line_lower:
            if _is_checked(line) and "grace" in line_lower:
                result["risk_level"] = "intermediate_low"

        # Shock / VT-VF detection
        if "shock" in line_lower and _is_checked(line):
            result["shock"] = True
        if ("vt/vf" in line_lower or "vt" in line_lower) and _is_checked(line):
            result["vt_vf"] = True

    return result


def _parse_cag_procedure(lines, full_text):
    """Section I: CAG Procedure."""
    result = {
        "cag_date": "",
        "pre_cag_risk": "",
        "physician_name": "",
        "physician_license": "",
        "access_site": "",
        "lesions": "",  # 1-vv/2-vv/3-vv/LM/non-obstructive
        "stenosis_locations": [],  # [{location, percent_range}]
        "complications": [],
    }

    in_cag_section = False

    for line in lines:
        line_lower = line.lower()

        if "cag procedure" in line_lower or "i. cag" in line_lower:
            in_cag_section = True

        if "pci procedure" in line_lower or "ii. pci" in line_lower:
            in_cag_section = False

        if not in_cag_section:
            continue

        # Date
        if "date of cag" in line_lower:
            m = re.search(r'(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})', line)
            if m:
                result["cag_date"] = m.group(1)

        # Pre-CAG Risk
        risk_options = ["very high", "high", "intermediate", "low"]
        for opt in risk_options:
            if opt in line_lower and _is_checked(line):
                result["pre_cag_risk"] = opt
                break

        # Physician
        if "responsible physician" in line_lower or "name" in line_lower:
            name = _extract_after(line, "Name")
            if name:
                result["physician_name"] = name
            lic = re.search(r'[วw]\.?\s*(\d+)', line)
            if lic:
                result["physician_license"] = lic.group(1)

        # Access site
        for site in ["femoral", "radial", "brachial"]:
            if site in line_lower and _is_checked(line):
                result["access_site"] = site

        # Lesions
        for les in ["1-vv", "2-vv", "3-vv", "LM", "non obstructive"]:
            if les.lower() in line_lower and _is_checked(line):
                result["lesions"] = les.replace(" ", "_")

        # Stenosis
        stenosis_ranges = ["1-25%", "26-50%", "51-75%", "76-90%", "91-99%", "100%"]
        for sr in stenosis_ranges:
            if sr in line and _is_checked(line):
                result["stenosis_locations"].append(sr)

        # Complications
        comp_list = ["death", "stroke", "ARF", "allergic", "hematoma",
                     "perforation", "dissection", "ACS", "shock", "arrest", "bleeding"]
        for comp in comp_list:
            if comp.lower() in line_lower and _is_checked(line):
                result["complications"].append(comp)

    return result


def _parse_pci_procedure(lines, full_text):
    """Section II: PCI Procedure."""
    result = {
        "pci_date": "",
        "access_site": "",
        "target_arteries": [],  # [{name, baseline_stenosis, lesion_type, syntax, stent_name, stent_type, final_stenosis, final_timi}]
        "syntax_score": "",  # low/intermediate/high
        "syntax_category": "",
        "complications": [],
    }

    in_pci_section = False
    current_artery = None

    artery_names = ["LM", "LAD/DG", "LAD", "Cx/OM", "Trifurcation", "RCA/PDA", "RCA"]

    for line in lines:
        line_lower = line.lower()

        if "pci procedure" in line_lower or "ii. pci" in line_lower:
            in_pci_section = True
        if not in_pci_section:
            continue

        # Date
        if "date of pci" in line_lower:
            m = re.search(r'(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})', line)
            if m:
                result["pci_date"] = m.group(1)

        # Target artery section
        for art in artery_names:
            if art.lower() in line_lower and ("segment" in line_lower or "target" in line_lower):
                current_artery = {"name": art, "baseline_stenosis": "", "lesion_type": "",
                                  "syntax": "", "stent_name": "", "stent_type": "",
                                  "final_stenosis": "", "pre_timi": "", "final_timi": ""}
                result["target_arteries"].append(current_artery)

        if current_artery:
            # Baseline stenosis
            if "baseline" in line_lower and "stenosis" in line_lower:
                for sr in ["1-25%", "26-50%", "51-75%", "76-90%", "91-99%", "100%"]:
                    if sr in line and _is_checked(line):
                        current_artery["baseline_stenosis"] = sr

            # Lesion type
            for lt in ["type A", "type B1", "type B2", "type C"]:
                if lt.lower() in line_lower and _is_checked(line):
                    current_artery["lesion_type"] = lt

            # SYNTAX score
            if "syntax" in line_lower:
                if "low" in line_lower and "(0-22)" in line and _is_checked(line):
                    result["syntax_score"] = "low (0-22)"
                    result["syntax_category"] = "low"
                elif "intermediate" in line_lower and "(23-32)" in line and _is_checked(line):
                    result["syntax_score"] = "intermediate (23-32)"
                    result["syntax_category"] = "intermediate"
                elif "high" in line_lower and "(>33)" in line and _is_checked(line):
                    result["syntax_score"] = "high (>33)"
                    result["syntax_category"] = "high"

            # Stent
            if "stent use" in line_lower or "stent" in line_lower:
                if "des" in line_lower and _is_checked(line):
                    current_artery["stent_type"] = "DES"
                elif "bms" in line_lower and _is_checked(line):
                    current_artery["stent_type"] = "BMS"
                # Stent name/size
                m = re.search(r'(?:specify|ชื่อ)[:\s]*(.+?)(?:\s+(?:BMS|DES)|\s*$)', line, re.IGNORECASE)
                if m:
                    name = re.sub(r'[…\.]+', '', m.group(1)).strip()
                    if name and len(name) > 2:
                        current_artery["stent_name"] = name

            # TIMI
            for timi in ["TIMI-0", "TIMI-1", "TIMI-2", "TIMI-3"]:
                if timi.lower() in line_lower and _is_checked(line):
                    if "pre" in line_lower:
                        current_artery["pre_timi"] = timi
                    elif "final" in line_lower:
                        current_artery["final_timi"] = timi

            # Final stenosis
            if "final" in line_lower and "stenosis" in line_lower:
                for sr in ["1-25%", "26-50%", "51-75%", "76-90%", "91-99%", "100%"]:
                    if sr in line and _is_checked(line):
                        current_artery["final_stenosis"] = sr

        # PCI Complications
        if "complication" in line_lower:
            comp_list = ["death", "stroke", "ARF", "hematoma", "perforation",
                         "dissection", "no reflow", "shock", "arrest"]
            for comp in comp_list:
                if comp.lower() in line_lower and _is_checked(line):
                    result["complications"].append(comp)

    return result


# ─── Appeal Data Extractor ──────────────────────────────

def cag_to_appeal_data(cag_data):
    """Convert parsed CAG form → dict for filling V993 appeal placeholders."""
    gen = cag_data["general"]
    risk = cag_data["risk_factors"]
    ind = cag_data["indication"]
    cag = cag_data["cag_procedure"]
    pci = cag_data["pci_procedure"]

    # Comorbidities list
    comorbidities = []
    if risk.get("dm"):
        comorbidities.append("DM")
    if risk.get("hypertension"):
        comorbidities.append("HT")
    if risk.get("dyslipidemia"):
        comorbidities.append("DLP")
    if risk.get("cerebrovascular_disease"):
        comorbidities.append("CVD")
    if risk.get("prior_mi"):
        comorbidities.append("Prior MI")

    # Stenosis summary
    stenosis_list = []
    for art in pci.get("target_arteries", []):
        if art["baseline_stenosis"]:
            stenosis_list.append(f'{art["name"]}: {art["baseline_stenosis"]}')

    # Stent summary
    stent_list = []
    for art in pci.get("target_arteries", []):
        if art["stent_name"] or art["stent_type"]:
            stent_list.append(f'{art["name"]}: {art["stent_name"]} ({art["stent_type"]})')

    return {
        "pid": gen.get("pid", ""),
        "sex": gen.get("sex", ""),
        "admission_date": gen.get("admission_date", ""),
        "case_type": ind.get("type", ""),
        "lvef": ind.get("lvef", ""),
        "risk_level": ind.get("risk_level", ""),
        "shock": ind.get("shock", False),
        "comorbidities": ", ".join(comorbidities) if comorbidities else "ไม่มี",
        "cag_date": cag.get("cag_date", ""),
        "physician": cag.get("physician_name", ""),
        "lesions": cag.get("lesions", ""),
        "access_site": cag.get("access_site", "") or pci.get("access_site", ""),
        "stenosis_summary": "; ".join(stenosis_list) if stenosis_list else "-",
        "stent_summary": "; ".join(stent_list) if stent_list else "-",
        "syntax_score": pci.get("syntax_score", ""),
        "syntax_category": pci.get("syntax_category", ""),
        "complications_cag": ", ".join(cag.get("complications", [])) or "ไม่มี",
        "complications_pci": ", ".join(pci.get("complications", [])) or "ไม่มี",
        "bmi": risk.get("bmi"),
    }


# ─── CLI ────────────────────────────────────────────────

if __name__ == "__main__":
    import sys, json

    if len(sys.argv) < 2:
        print("Usage: python scripts/cag_form_parser.py <filled_cag_form.docx>")
        sys.exit(1)

    path = sys.argv[1]
    data = parse_cag_form(path)
    appeal = cag_to_appeal_data(data)

    print(json.dumps(data, ensure_ascii=False, indent=2, default=str))
    print("\n--- Appeal Data ---")
    print(json.dumps(appeal, ensure_ascii=False, indent=2, default=str))
